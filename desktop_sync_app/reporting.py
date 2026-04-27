from __future__ import annotations

import json
from collections import Counter, defaultdict
import copy
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from text_polish import GrantProofTextPolisher
from premium_report import PremiumActivityReportBuilder
from premium_excel_report import PremiumExcelActivityReportBuilder
from premium_pptx_report import PremiumPptxActivityReportBuilder

IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.webp'}
VIDEO_EXTENSIONS = {'.mp4', '.mov', '.m4v', '.avi', '.mkv', '.webm'}

BLUE = '173B72'
ORANGE = 'F28C28'
SOFT_BLUE = 'EAF1FB'
SOFT_ORANGE = 'FFF2E2'
SOFT_GREEN = 'ECFDF3'
SOFT_GREY = 'F6F8FC'
MID_GREY = '6B7280'
DARK = '1F2937'

THIN_GREY_BORDER = Border(
    left=Side(style='thin', color='D1D9E6'),
    right=Side(style='thin', color='D1D9E6'),
    top=Side(style='thin', color='D1D9E6'),
    bottom=Side(style='thin', color='D1D9E6'),
)


def _safe_text(value: object, fallback: str = '') -> str:
    if value is None:
        return fallback
    text = str(value).strip()
    return text if text else fallback


def _safe_date(value: object) -> str:
    text = _safe_text(value)
    if not text:
        return ''
    try:
        return datetime.fromisoformat(text.replace('Z', '+00:00')).strftime('%Y-%m-%d %H:%M')
    except Exception:
        return text


def _parse_iso(value: object) -> datetime | None:
    text = _safe_text(value)
    if not text:
        return None
    try:
        return datetime.fromisoformat(text.replace('Z', '+00:00'))
    except Exception:
        return None


def _title_case(value: str) -> str:
    cleaned = _safe_text(value)
    return cleaned[:1].upper() + cleaned[1:] if cleaned else ''


def _sentence(value: str, fallback: str = '') -> str:
    text = ' '.join(_safe_text(value, fallback).replace('\n', ' ').split())
    if not text:
        return fallback
    if text[-1] not in '.!?':
        text += '.'
    return text


def _comma_join(items: list[str], fallback: str = '-') -> str:
    clean = [item for item in (_safe_text(i) for i in items) if item]
    return ', '.join(clean) if clean else fallback


def _normalize_lang(value: object) -> str:
    text = _safe_text(value).lower()
    if text.startswith('en'):
        return 'en'
    return 'fr'


@dataclass
class ItemRecord:
    kind: str
    subtype: str
    project_code: str
    project_name: str
    title: str
    description: str
    created_at: str
    location: str
    relative_folder: str
    metadata_path: Path
    media_files: list[Path]
    raw: dict


class GrantProofAIWriter:
    """Local, deterministic writing layer for donor-friendly narratives."""

    def executive_summary(self, project_name: str, donor_name: str, country: str, records: list[ItemRecord], lang: str) -> list[str]:
        evidence_records = [record for record in records if record.kind == 'evidence']
        story_records = [record for record in records if record.kind == 'story']
        outputs = [self._output_label(record) for record in evidence_records if self._output_label(record)]
        activities = [self._activity_label(record) for record in evidence_records if self._activity_label(record)]
        locations = [self._location_label(record) for record in records if self._location_label(record)]
        media_count = sum(len(record.media_files) for record in records)
        gps_count = sum(1 for record in evidence_records if record.raw.get('latitude') is not None and record.raw.get('longitude') is not None)
        top_output = self._top_label(outputs, 'les outputs attendus du projet' if lang == 'fr' else 'the expected project outputs')
        top_activity = self._top_label(activities, 'les activités de mise en œuvre' if lang == 'fr' else 'implementation activities')
        top_location = self._top_label(locations, 'les sites d’intervention documentés' if lang == 'fr' else 'documented field locations')
        descriptive_signals = self._description_signals(records)
        implementation_maturity = self._implementation_maturity(records, lang)

        if lang == 'fr':
            donor_clause = f' financé par {donor_name}' if donor_name else ''
            country_clause = f' au {country}' if country else ''
            para1 = (
                f'Ce dossier GrantProof consolide les dernières preuves terrain collectées pour {project_name}{country_clause}{donor_clause}. '
                f'La période actuelle rassemble {len(evidence_records)} preuve(s), {len(story_records)} story(s) et {media_count} média(s) synchronisés, '
                f'créant une base directement réutilisable pour des rapports narratifs, des annexes illustrées et des revues bailleurs à court délai.'
            )
            para2 = (
                f'Le signal de mise en œuvre le plus solide concerne {top_activity}, avec une contribution particulièrement visible autour de {top_output}. '
                f'Le portefeuille documenté montre une continuité crédible entre activité conduite, résultats immédiats observables et éléments de vérification visuelle ou qualitative.'
            )
            para3 = (
                f'L’empreinte opérationnelle la plus lisible ressort à {top_location}. '
                f'Les textes synchronisés font ressortir {descriptive_signals}, ce qui donne au dossier un niveau de densité suffisant pour soutenir un reporting donateurs plus convaincant et moins dépendant d’une réécriture manuelle.'
            )
            para4 = (
                f'La maturité documentaire actuelle peut être qualifiée de {implementation_maturity}. '
                f'Elle est renforcée par {gps_count} élément(s) géolocalisé(s) ainsi que par une base média assez structurée pour produire des annexes propres, cohérentes et immédiatement partageables.'
            )
        else:
            donor_clause = f' funded by {donor_name}' if donor_name else ''
            country_clause = f' in {country}' if country else ''
            para1 = (
                f'This GrantProof pack consolidates the latest field evidence captured for {project_name}{country_clause}{donor_clause}. '
                f'The current reporting window brings together {len(evidence_records)} evidence item(s), {len(story_records)} story item(s), and {media_count} synced media asset(s), '
                f'creating a base that can be reused immediately for narrative reporting, illustrated annexes, and donor review packs.'
            )
            para2 = (
                f'The strongest implementation signal relates to {top_activity}, with the clearest documented contribution linked to {top_output}. '
                f'The documented portfolio shows a credible continuity between delivered activities, observable short-term results, and supporting visual or qualitative verification.'
            )
            para3 = (
                f'The clearest operational footprint appears in {top_location}. '
                f'Synced descriptions consistently point to {descriptive_signals}, giving the pack enough density to support stronger donor narratives with less manual rewriting.'
            )
            para4 = (
                f'The current level of documentation can be described as {implementation_maturity}. '
                f'It is reinforced by {gps_count} geo-referenced item(s) and by a media base that is already structured enough to produce clean, coherent, and ready-to-share annexes.'
            )
        return [_sentence(para1), _sentence(para2), _sentence(para3), _sentence(para4)]

    def key_highlights(self, records: list[ItemRecord], lang: str) -> list[str]:
        evidence_records = [record for record in records if record.kind == 'evidence']
        story_records = [record for record in records if record.kind == 'story']
        activity_counts = Counter(self._activity_label(record) for record in evidence_records if self._activity_label(record))
        output_counts = Counter(self._output_label(record) for record in evidence_records if self._output_label(record))
        media_rich = sum(1 for record in records if record.media_files)
        gps_enabled = sum(1 for record in evidence_records if record.raw.get('latitude') is not None and record.raw.get('longitude') is not None)
        with_descriptions = sum(1 for record in records if _safe_text(record.description))
        coverage_ratio = int((with_descriptions / len(records)) * 100) if records else 0

        if lang == 'fr':
            highlights = [
                f'{len(evidence_records)} preuve(s) et {len(story_records)} story(s) sont déjà packagées pour un usage direct dans le reporting bailleur.',
                f'{media_rich} élément(s) disposent de média(s) liés, ce qui renforce la valeur de vérification et la qualité des annexes.',
                f'{with_descriptions} élément(s) contiennent un texte descriptif exploitable, soit une couverture narrative d’environ {coverage_ratio} % du portefeuille synchronisé.',
            ]
        else:
            highlights = [
                f'{len(evidence_records)} evidence item(s) and {len(story_records)} story item(s) are already packaged for direct reuse in donor reporting.',
                f'{media_rich} captured item(s) include linked media, significantly strengthening verification value and annex quality.',
                f'{with_descriptions} item(s) contain usable descriptive text, representing narrative coverage of roughly {coverage_ratio}% of the synced portfolio.',
            ]
        if activity_counts:
            activity, count = activity_counts.most_common(1)[0]
            highlights.append(
                f"Le cluster d’activité le plus documenté est {activity} avec {count} preuve(s) liées." if lang == 'fr' else f'The most documented activity cluster is {activity} with {count} linked evidence item(s).'
            )
        if output_counts:
            output, count = output_counts.most_common(1)[0]
            highlights.append(
                f"Le signal output le plus fort concerne {output} avec {count} preuve(s) liées." if lang == 'fr' else f'The strongest documented output signal is {output} with {count} linked evidence item(s).'
            )
        if gps_enabled:
            highlights.append(
                f'{gps_enabled} preuve(s) incluent des coordonnées GPS, ce qui renforce la traçabilité pour la vérification et l’audit.'
                if lang == 'fr'
                else f'{gps_enabled} evidence item(s) include GPS coordinates, strengthening traceability for verification and audit purposes.'
            )
        return [_sentence(item) for item in highlights[:6]]

    def implementation_summary(self, activity: str, records: list[ItemRecord], lang: str) -> str:
        outputs = [self._output_label(record) for record in records if self._output_label(record)]
        locations = [self._location_label(record) for record in records if self._location_label(record)]
        media_count = sum(len(record.media_files) for record in records)
        top_output = self._top_label(outputs, 'les outputs liés' if lang == 'fr' else 'linked outputs')
        top_location = self._top_label(locations, 'les sites documentés' if lang == 'fr' else 'documented field locations')
        evidence_titles = ', '.join(record.title for record in records[:3])
        if lang == 'fr':
            return _sentence(
                f'Le cluster d’activité « {activity} » regroupe {len(records)} élément(s) terrain et {media_count} média(s) associés. '
                f'Le matériau disponible renvoie principalement à {top_output} et montre une empreinte de mise en œuvre particulièrement lisible à {top_location}. '
                f'Les pièces les plus représentatives portent notamment sur {evidence_titles or "les éléments synchronisés"}, ce qui permet de rédiger un paragraphe bailleur plus dense, plus crédible et mieux relié aux résultats observables.'
            )
        return _sentence(
            f'The activity cluster “{activity}” brings together {len(records)} field item(s) and {media_count} supporting media asset(s). '
            f'The available material points most clearly to {top_output} and shows a particularly visible implementation footprint in {top_location}. '
            f'The most representative entries include {evidence_titles or "the synced items"}, which provides a stronger base for denser and more credible donor-facing drafting.'
        )

    def activity_deep_dive(self, activity: str, records: list[ItemRecord], lang: str) -> list[str]:
        outputs = sorted({self._output_label(record) for record in records if self._output_label(record)})
        locations = sorted({self._location_label(record) for record in records if self._location_label(record)})
        media_count = sum(len(record.media_files) for record in records)
        description_fragments = [record.description for record in records if _safe_text(record.description)]
        description_hint = '; '.join(description_fragments[:2]) if description_fragments else ('des signaux de mise en œuvre observables' if lang == 'fr' else 'observable implementation signals')
        if lang == 'fr':
            return [
                _sentence(f'Pour l’activité « {activity} », GrantProof consolide {len(records)} élément(s) documentés, adossés à {media_count} média(s), ce qui donne un socle suffisamment robuste pour une restitution bailleur structurée.'),
                _sentence(f'Les outputs les plus directement reliés à cette activité sont {_comma_join(outputs)}, avec une concentration d’observations à {_comma_join(locations)}.'),
                _sentence(f'Les descriptions synchronisées mettent surtout en avant {description_hint}, ce qui permet de transformer des notes terrain courtes en paragraphe de reporting plus solide et plus institutionnel.'),
            ]
        return [
            _sentence(f'For the activity “{activity}”, GrantProof consolidates {len(records)} documented item(s) supported by {media_count} media asset(s), creating a sufficiently robust base for structured donor reporting.'),
            _sentence(f'The outputs most directly linked to this activity are {_comma_join(outputs)}, with observations concentrated in {_comma_join(locations)}.'),
            _sentence(f'The synced descriptions primarily point to {description_hint}, making it easier to convert short field notes into stronger and more institutional reporting language.'),
        ]

    def recommended_next_steps(self, records: list[ItemRecord], lang: str) -> list[str]:
        evidence_records = [record for record in records if record.kind == 'evidence']
        story_records = [record for record in records if record.kind == 'story']
        without_media = sum(1 for record in records if not record.media_files)
        without_description = sum(1 for record in records if not _safe_text(record.description))
        without_gps = sum(1 for record in evidence_records if record.raw.get('latitude') is None or record.raw.get('longitude') is None)
        next_steps: list[str] = []
        if lang == 'fr':
            if without_description:
                next_steps.append(f'Compléter les descriptions pour {without_description} élément(s) améliorerait nettement la densité narrative du rapport final.')
            if without_media:
                next_steps.append(f'L’ajout de médias complémentaires sur {without_media} entrée(s) renforcerait la valeur de preuve et la qualité des annexes.')
            if without_gps:
                next_steps.append(f'La capture GPS de {without_gps} preuve(s) supplémentaires améliorerait la traçabilité spatiale du portefeuille.')
            if story_records:
                next_steps.append('Un meilleur usage des stories et citations bénéficiaires permettrait encore de renforcer la dimension qualitative du dossier.')
            if not next_steps:
                next_steps.append('Le niveau de complétude actuel permet déjà un usage très correct pour le reporting bailleur et les revues de projet.')
        else:
            if without_description:
                next_steps.append(f'Completing the descriptions for {without_description} item(s) would significantly improve the narrative density of the final report.')
            if without_media:
                next_steps.append(f'Adding supporting media to {without_media} entry(ies) would strengthen evidentiary value and annex quality.')
            if without_gps:
                next_steps.append(f'Capturing GPS coordinates for {without_gps} additional evidence item(s) would improve the spatial traceability of the portfolio.')
            if story_records:
                next_steps.append('Making fuller use of stories and beneficiary quotes would further strengthen the qualitative dimension of the pack.')
            if not next_steps:
                next_steps.append('The current level of completeness is already strong enough for donor reporting and project review use.')
        return [_sentence(item) for item in next_steps[:4]]

    def _implementation_maturity(self, records: list[ItemRecord], lang: str) -> str:
        if not records:
            return 'faible' if lang == 'fr' else 'limited'
        media_count = sum(len(record.media_files) for record in records)
        described = sum(1 for record in records if _safe_text(record.description))
        stories = sum(1 for record in records if record.kind == 'story')
        score = media_count + described + stories
        if lang == 'fr':
            if score >= max(8, len(records) * 2):
                return 'solide'
            if score >= max(4, len(records)):
                return 'intermédiaire'
            return 'encore perfectible'
        if score >= max(8, len(records) * 2):
            return 'strong'
        if score >= max(4, len(records)):
            return 'moderately strong'
        return 'still developing'

    def evidence_narrative(self, record: ItemRecord, lang: str) -> str:
        activity = self._activity_label(record) or ('le flux d’activité prévu' if lang == 'fr' else 'the planned activity stream')
        output = self._output_label(record) or ('l’output attendu' if lang == 'fr' else 'the intended output')
        location = self._location_label(record) or ('le site documenté' if lang == 'fr' else 'the documented field site')
        description = _safe_text(record.description)
        subtitle = self._subtype_phrase(record, lang)
        media_count = len(record.media_files)
        media_clause = (
            f' et s’appuie sur {media_count} média(s) lié(s)' if lang == 'fr' else f' and is supported by {media_count} linked media asset(s)'
        ) if media_count else ''
        if description:
            if lang == 'fr':
                return _sentence(
                    f'{subtitle} capturée dans le cadre de {activity} à {location}. {description} '
                    f'Pour un usage bailleur, cette entrée apporte une trace concrète de mise en œuvre reliée à {output}{media_clause}, tout en fournissant une matière immédiatement mobilisable pour illustrer des résultats observés sans surinterprétation.'
                )
            return _sentence(
                f'{subtitle} captured under {activity} in {location}. {description} '
                f'For donor-facing use, this entry provides a concrete implementation trace linked to {output}{media_clause}, while also supplying material that can be reused immediately to illustrate observed results without overstating the evidence.'
            )
        if lang == 'fr':
            return _sentence(
                f'{subtitle} capturée dans le cadre de {activity} à {location}. '
                f'Même avec peu de texte libre, cet élément constitue un signal vérifiable de progression vers {output}{media_clause} et reste utile pour documenter la réalité opérationnelle du terrain.'
            )
        return _sentence(
            f'{subtitle} captured under {activity} in {location}. '
            f'Even with limited free-text detail, this item still provides a verifiable signal of progress toward {output}{media_clause} and remains useful for documenting operational reality.'
        )

    def story_narrative(self, record: ItemRecord, lang: str) -> str:
        beneficiary = self._location_label(record) or ('la personne documentée' if lang == 'fr' else 'the documented participant')
        summary = _safe_text(record.description)
        quote = _safe_text(record.raw.get('quote'))
        quote_clause = (
            f' La citation enregistrée — « {quote} » — apporte une voix directe qui renforce la crédibilité qualitative du changement observé.'
            if lang == 'fr' and quote else
            f' The recorded quote — “{quote}” — adds a direct voice that strengthens the qualitative credibility of the observed change.' if quote else ''
        )
        if summary:
            if lang == 'fr':
                return _sentence(
                    f'Cette story met en lumière {beneficiary}. {summary} '
                    f'Pour le reporting bailleur, elle complète utilement la trace opérationnelle en donnant une lecture plus humaine, plus contextualisée et plus mobilisable des résultats observés.{quote_clause}'
                )
            return _sentence(
                f'This story focuses on {beneficiary}. {summary} '
                f'For donor reporting, it usefully complements the operational trail by offering a more human, contextualized, and reusable reading of the observed results.{quote_clause}'
            )
        if lang == 'fr':
            return _sentence(
                f'Cette story met en lumière {beneficiary}. '
                f'Elle apporte une couche qualitative à la base de preuves et aide à transformer la mise en œuvre du projet en récit plus incarné.{quote_clause}'
            )
        return _sentence(
            f'This story focuses on {beneficiary}. '
            f'It adds a qualitative layer to the evidence base and helps turn project implementation into a more human-centered narrative.{quote_clause}'
        )

    def item_takeaways(self, record: ItemRecord, lang: str) -> list[str]:
        bullets: list[str] = []
        if record.kind == 'evidence':
            activity = self._activity_label(record)
            output = self._output_label(record)
            if activity:
                bullets.append(f'Activité liée : {activity}.' if lang == 'fr' else f'Linked activity: {activity}.')
            if output:
                bullets.append(f'Output illustré : {output}.' if lang == 'fr' else f'Illustrated output: {output}.')
            if record.raw.get('latitude') is not None and record.raw.get('longitude') is not None:
                bullets.append('Des coordonnées GPS sont disponibles pour la traçabilité.' if lang == 'fr' else 'GPS coordinates are available for traceability.')
        else:
            beneficiary = self._location_label(record)
            if beneficiary:
                bullets.append(f'Perspective bénéficiaire documentée : {beneficiary}.' if lang == 'fr' else f'Beneficiary perspective documented: {beneficiary}.')
            if record.raw.get('consentGiven'):
                bullets.append('Le consentement a été confirmé au moment de la collecte.' if lang == 'fr' else 'Consent was confirmed at the time of capture.')
        if record.media_files:
            bullets.append(
                f'{len(record.media_files)} média(s) soutiennent cette entrée.' if lang == 'fr' else f'{len(record.media_files)} media asset(s) support this entry.'
            )
        if not bullets:
            bullets.append('Cet élément reste disponible comme entrée de reporting documentée.' if lang == 'fr' else 'This item remains available as a documented reporting input.')
        return [_sentence(item) for item in bullets]

    def media_caption(self, record: ItemRecord, file_path: Path, lang: str) -> str:
        base = file_path.stem.replace('_', ' ').replace('-', ' ').strip() or file_path.name
        if file_path.suffix.lower() in IMAGE_EXTENSIONS:
            return _sentence(
                f'Image illustrative liée à {record.title} : {base}.' if lang == 'fr' else f'Illustrative image linked to {record.title}: {base}.'
            )
        if file_path.suffix.lower() in VIDEO_EXTENSIONS:
            return _sentence(
                f'Vidéo liée à {record.title} : {base}.' if lang == 'fr' else f'Video evidence linked to {record.title}: {base}.'
            )
        return _sentence(
            f'Fichier de support lié à {record.title} : {base}.' if lang == 'fr' else f'Linked supporting file for {record.title}: {base}.'
        )

    def monitoring_note(self, records: list[ItemRecord], lang: str) -> list[str]:
        evidence_records = [record for record in records if record.kind == 'evidence']
        story_records = [record for record in records if record.kind == 'story']
        gps_count = sum(1 for record in evidence_records if record.raw.get('latitude') is not None and record.raw.get('longitude') is not None)
        consent_count = sum(1 for record in story_records if record.raw.get('consentGiven'))
        media_count = sum(len(record.media_files) for record in records)
        if lang == 'fr':
            return [
                _sentence(f'Le pack repose sur {media_count} média(s) synchronisé(s), ce qui offre une base robuste pour les annexes visuelles, la communication de résultats et la vérification croisée.'),
                _sentence(f'{gps_count} preuve(s) incluent des coordonnées GPS, renforçant la traçabilité spatiale du dossier pour la revue et l’audit.'),
                _sentence(f'{consent_count} story(s) indiquent un consentement confirmé, ce qui soutient une réutilisation plus sûre des récits qualitatifs et des citations directes.'),
            ]
        return [
            _sentence(f'The pack relies on {media_count} synced media asset(s), providing a robust base for visual annexes, visibility material, and cross-verification.'),
            _sentence(f'{gps_count} evidence item(s) include GPS coordinates, strengthening the spatial traceability of the pack for review and audit purposes.'),
            _sentence(f'{consent_count} story item(s) indicate confirmed consent, supporting safer reuse of qualitative narratives and direct quotations.'),
        ]

    def _activity_label(self, record: ItemRecord) -> str:
        return _safe_text(record.raw.get('activity'))

    def _output_label(self, record: ItemRecord) -> str:
        return _safe_text(record.raw.get('output'))

    def _location_label(self, record: ItemRecord) -> str:
        return _safe_text(record.raw.get('locationLabel') or record.raw.get('beneficiaryAlias'))

    def _top_label(self, values: list[str], fallback: str) -> str:
        clean = [value for value in values if value]
        if not clean:
            return fallback
        return Counter(clean).most_common(1)[0][0]

    def _subtype_phrase(self, record: ItemRecord, lang: str) -> str:
        subtype = _safe_text(record.subtype)
        mapping_fr = {
            'attendance': 'Preuve de présence',
            'document': 'Preuve documentaire',
            'receipt': 'Justificatif ou reçu',
            'note': 'Note de terrain',
            'photo': 'Preuve photo',
            'video': 'Preuve vidéo',
            'story': 'Story terrain',
        }
        mapping_en = {
            'attendance': 'Attendance evidence',
            'document': 'Documentary evidence',
            'receipt': 'Supporting receipt or justification',
            'note': 'Field note',
            'photo': 'Photo evidence',
            'video': 'Video evidence',
            'story': 'Field story',
        }
        if lang == 'fr':
            return mapping_fr.get(subtype, 'Preuve terrain')
        return mapping_en.get(subtype, 'Field evidence')

    def _description_signals(self, records: list[ItemRecord]) -> str:
        fragments: list[str] = []
        for record in records:
            description = _safe_text(record.description)
            if description:
                fragments.append(description[:120].strip())
        if not fragments:
            return 'des observations de terrain directement liées à la mise en œuvre'
        top = fragments[:3]
        if len(top) == 1:
            return top[0]
        return '; '.join(top)


class GrantProofReportEngine:
    def __init__(self, base_folder: Path, organization_name: str = '') -> None:
        self.base_folder = Path(base_folder)
        self.organization_name = _safe_text(organization_name)
        self.projects_root = self.base_folder / 'projects'
        self.ai_writer = GrantProofAIWriter()
        self.text_polisher = GrantProofTextPolisher()
        self.premium_report_builder = PremiumActivityReportBuilder(self.base_folder, default_org_name=self.organization_name)
        self.premium_excel_report_builder = PremiumExcelActivityReportBuilder(self.base_folder, default_org_name=self.organization_name)
        self.premium_pptx_report_builder = PremiumPptxActivityReportBuilder(self.base_folder, default_org_name=self.organization_name)

    def ensure_root_files(self) -> None:
        self.base_folder.mkdir(parents=True, exist_ok=True)
        readme = self.base_folder / '_README_GrantProof.txt'
        readme.write_text(
            '\n'.join([
                'GrantProof export folder',
                '',
                'This folder is maintained by GrantProof Desktop Sync.',
                'Projects contain evidence, stories and ready-to-share report outputs generated from mobile sync.',
                'Premium reports are available in each project under the reports folder.',
                '',
                'Main outputs:',
                '- Project_Report_FR.pptx / Project_Report_EN.pptx',
                '- Project_Register_FR.xlsx / Project_Register_EN.xlsx',
                '',
                'Technical data remains under projects/<code>/evidence and projects/<code>/stories.',
                'Those folders store JSON and media only; legacy evidence.docx/story.docx files are removed during rebuild.',
            ]),
            encoding='utf-8',
        )
        self.rebuild_global_index()

    def rebuild_global_index(self) -> None:
        records = self._collect_all_records()
        wb = Workbook()
        ws = wb.active
        ws.title = 'Projects'
        ws.freeze_panes = 'A2'
        ws.append(['Project code', 'Project name', 'Evidence', 'Stories', 'Media assets', 'Last update'])
        self._header_style(ws[1], fill=BLUE)

        projects: dict[str, dict[str, object]] = {}
        for record in records:
            stats = projects.setdefault(record.project_code, {
                'name': record.project_name,
                'evidence': 0,
                'stories': 0,
                'media': 0,
                'last': '',
                'lang': self._record_language(record),
            })
            if record.kind == 'evidence':
                stats['evidence'] = int(stats['evidence']) + 1
            else:
                stats['stories'] = int(stats['stories']) + 1
            stats['media'] = int(stats['media']) + len(record.media_files)
            last_value = _safe_date(record.raw.get('syncedAt') or record.raw.get('createdAt'))
            if last_value and (not stats['last'] or str(last_value) > str(stats['last'])):
                stats['last'] = last_value
                stats['lang'] = self._record_language(record)

        for code in sorted(projects.keys()):
            stats = projects[code]
            ws.append([code, stats['name'], stats['evidence'], stats['stories'], stats['media'], stats['last']])

        details = wb.create_sheet('Items')
        details.freeze_panes = 'A2'
        details.append(['Kind', 'Subtype', 'Project code', 'Project name', 'Title', 'Created at', 'Location', 'Media', 'Report lang', 'Folder'])
        self._header_style(details[1], fill=ORANGE)
        for record in records:
            details.append([
                record.kind,
                record.subtype,
                record.project_code,
                record.project_name,
                record.title,
                record.created_at,
                record.location,
                len(record.media_files),
                self._record_language(record).upper(),
                record.relative_folder,
            ])

        summary = wb.create_sheet('Summary')
        summary.append(['Metric', 'Value'])
        self._header_style(summary[1], fill=BLUE)
        summary.append(['Projects tracked', len(projects)])
        summary.append(['Evidence total', sum(int(item['evidence']) for item in projects.values())])
        summary.append(['Stories total', sum(int(item['stories']) for item in projects.values())])
        summary.append(['Media total', sum(int(item['media']) for item in projects.values())])
        summary.append(['Generated at', datetime.now().strftime('%Y-%m-%d %H:%M')])

        for sheet in (ws, details, summary):
            self._autosize(sheet)
            self._apply_sheet_polish(sheet, auto_filter=True)

        wb.save(self.base_folder / '_INDEX_GrantProof.xlsx')

    def rebuild_project(self, project_code: str, preferred_language: str | None = None) -> None:
        project_dir = self.projects_root / project_code
        if not project_dir.exists():
            return
        records = self._collect_project_records(project_code)
        if not records:
            return

        reports_dir = project_dir / 'reports'
        reports_dir.mkdir(parents=True, exist_ok=True)
        project_name = next((r.project_name for r in records if r.project_name), project_code)
        report_lang = self._detect_project_language(records, preferred_language)

        # Keep user-facing outputs in the project reports folder only.
        # The evidence/stories folders remain technical storage for JSON and media, but no longer receive Word reports.
        self._cleanup_legacy_report_outputs(project_dir, reports_dir)
        self._build_project_register(reports_dir / 'Project_Register_FR.xlsx', project_code, project_name, records, 'fr')
        self._build_project_register(reports_dir / 'Project_Register_EN.xlsx', project_code, project_name, records, 'en')
        self._build_project_pptx_report(reports_dir / 'Project_Report_FR.pptx', project_code, project_name, records, 'fr')
        self._build_project_pptx_report(reports_dir / 'Project_Report_EN.pptx', project_code, project_name, records, 'en')

        self.rebuild_global_index()

    def rebuild_all(self, preferred_language: str | None = None) -> None:
        self.ensure_root_files()
        if not self.projects_root.exists():
            return
        for project_dir in self.projects_root.iterdir():
            if project_dir.is_dir():
                self.rebuild_project(project_dir.name, preferred_language=preferred_language)

    def _cleanup_legacy_report_outputs(self, project_dir: Path, reports_dir: Path) -> None:
        # Remove older duplicate files so the user sees a clean reports folder.
        for legacy_name in ('Project_Report.docx', 'Project_Report_FR.docx', 'Project_Report_EN.docx', 'Project_Report_FR.xlsx', 'Project_Report_EN.xlsx', 'Project_Register.xlsx'):
            legacy = reports_dir / legacy_name
            if legacy.exists():
                try:
                    legacy.unlink()
                except Exception:
                    pass
        # Remove old per-item Word reports from technical data folders.
        for legacy_name in ('evidence.docx', 'story.docx'):
            for legacy in project_dir.rglob(legacy_name):
                try:
                    legacy.unlink()
                except Exception:
                    pass

    def _collect_all_records(self) -> list[ItemRecord]:
        records: list[ItemRecord] = []
        if not self.projects_root.exists():
            return records
        for project_dir in self.projects_root.iterdir():
            if project_dir.is_dir():
                records.extend(self._collect_project_records(project_dir.name))
        return sorted(records, key=lambda item: (item.project_code, item.created_at, item.title))

    def _collect_project_records(self, project_code: str) -> list[ItemRecord]:
        project_dir = self.projects_root / project_code
        if not project_dir.exists():
            return []
        records: list[ItemRecord] = []
        for kind in ('evidence', 'stories'):
            kind_dir = project_dir / kind
            if not kind_dir.exists():
                continue
            for metadata_path in kind_dir.rglob('*.json'):
                if metadata_path.name not in {'evidence.json', 'story.json'}:
                    continue
                try:
                    raw = json.loads(metadata_path.read_text(encoding='utf-8'))
                except Exception:
                    continue
                item_kind = 'evidence' if kind == 'evidence' else 'story'
                record_lang = _normalize_lang(raw.get('reportLanguage') or raw.get('languageCode') or 'fr')
                raw = self._polish_record_payload(raw, item_kind, project_code, record_lang)
                project = raw.get('project') or {}
                created_at = _safe_date(raw.get('createdAt'))
                title = _safe_text(raw.get('title'), metadata_path.parent.name)
                if kind == 'evidence':
                    description = _safe_text(raw.get('description'), '')
                    location = _safe_text(raw.get('locationLabel'), '-')
                    subtype = _safe_text(raw.get('type'), 'note')
                else:
                    description = _safe_text(raw.get('summary'), '')
                    location = _safe_text(raw.get('beneficiaryAlias') or raw.get('locationLabel'), '-')
                    subtype = 'story'
                media_dir = metadata_path.parent / 'media'
                media_files = sorted([path for path in media_dir.iterdir() if path.is_file()]) if media_dir.exists() else []
                records.append(
                    ItemRecord(
                        kind=item_kind,
                        subtype=subtype,
                        project_code=project_code,
                        project_name=_safe_text(project.get('name'), project_code),
                        title=title,
                        description=description,
                        created_at=created_at,
                        location=location,
                        relative_folder=str(metadata_path.parent.relative_to(self.base_folder)),
                        metadata_path=metadata_path,
                        media_files=media_files,
                        raw=raw,
                    )
                )
        return sorted(records, key=lambda item: (item.created_at, item.title))

    def _polish_project_payload(self, project: dict, lang: str) -> dict:
        polished = copy.deepcopy(project) if isinstance(project, dict) else {}
        polished['name'] = self.text_polisher.normalize_label(polished.get('name'), lang=lang, title_mode=True)
        polished['donorName'] = self.text_polisher.normalize_label(polished.get('donorName'), lang=lang, title_mode=True)
        polished['country'] = self.text_polisher.normalize_country(polished.get('country'), lang=lang)
        polished['activities'] = [
            self.text_polisher.normalize_label(item, lang=lang, title_mode=True)
            for item in polished.get('activities', [])
            if _safe_text(item)
        ]
        polished['outputs'] = [
            self.text_polisher.normalize_label(item, lang=lang, title_mode=True)
            for item in polished.get('outputs', [])
            if _safe_text(item)
        ]
        return polished

    def _polish_record_payload(self, raw: dict, item_kind: str, project_code: str, lang: str) -> dict:
        polished = copy.deepcopy(raw)
        project = self._polish_project_payload(polished.get('project') if isinstance(polished.get('project'), dict) else {}, lang)
        if project:
            polished['project'] = project
        activity_choices = project.get('activities', []) if isinstance(project.get('activities'), list) else []
        output_choices = project.get('outputs', []) if isinstance(project.get('outputs'), list) else []

        polished['title'] = self.text_polisher.normalize_label(polished.get('title'), lang=lang, title_mode=True)
        polished['description'] = self.text_polisher.normalize_sentence(polished.get('description'), lang=lang)
        polished['summary'] = self.text_polisher.normalize_sentence(polished.get('summary'), lang=lang)
        polished['locationLabel'] = self.text_polisher.normalize_label(polished.get('locationLabel'), lang=lang, title_mode=True)
        polished['beneficiaryAlias'] = self.text_polisher.normalize_label(polished.get('beneficiaryAlias'), lang=lang, title_mode=True)
        polished['quote'] = self.text_polisher.normalize_quote(polished.get('quote'), lang=lang)
        polished['activity'] = self.text_polisher.normalize_label(polished.get('activity'), lang=lang, choices=activity_choices, title_mode=True)
        polished['output'] = self.text_polisher.normalize_label(polished.get('output'), lang=lang, choices=output_choices, title_mode=True)
        if item_kind == 'story' and not polished.get('locationLabel'):
            polished['locationLabel'] = polished.get('beneficiaryAlias', '')
        return polished

    def _build_project_register(self, output: Path, project_code: str, project_name: str, records: list[ItemRecord], lang: str) -> None:
        project_meta = self._project_meta(records, preferred_language=lang)
        wb = Workbook()
        overview = wb.active
        overview.title = 'Vue' if lang == 'fr' else 'Overview'
        overview.append([self._label('metric', lang), self._label('value', lang)])
        self._header_style(overview[1], fill=BLUE)
        overview_rows = [
            (self._label('project_code', lang), project_code),
            (self._label('project_name', lang), project_name),
            (self._label('donor', lang), project_meta['donor_name']),
            (self._label('country', lang), project_meta['country']),
            (self._label('reporting_window', lang), project_meta['reporting_window']),
            (self._label('report_language', lang), self._label('language_french', lang) if lang == 'fr' else self._label('language_english', lang)),
            (self._label('generated_at', lang), datetime.now().strftime('%Y-%m-%d %H:%M')),
            (self._label('evidence_items', lang), sum(1 for r in records if r.kind == 'evidence')),
            (self._label('story_items', lang), sum(1 for r in records if r.kind == 'story')),
            (self._label('media_assets', lang), sum(len(r.media_files) for r in records)),
            (self._label('top_activity', lang), self._top_value([_safe_text(r.raw.get('activity')) for r in records if r.kind == 'evidence'], self._label('not_specified', lang))),
            (self._label('top_output', lang), self._top_value([_safe_text(r.raw.get('output')) for r in records if r.kind == 'evidence'], self._label('not_specified', lang))),
        ]
        for row in overview_rows:
            overview.append(list(row))

        narrative = wb.create_sheet('Narratif' if lang == 'fr' else 'Narrative')
        narrative.append([self._label('section', lang), self._label('content', lang)])
        self._header_style(narrative[1], fill=ORANGE)
        for index, paragraph in enumerate(self.ai_writer.executive_summary(project_name, project_meta['donor_name'], project_meta['country'], records, lang), start=1):
            narrative.append([f'{self._label("executive_summary", lang)} {index}', paragraph])
        for index, bullet in enumerate(self.ai_writer.key_highlights(records, lang), start=1):
            narrative.append([f'{self._label("key_highlights", lang)} {index}', bullet])
        for index, paragraph in enumerate(self.ai_writer.monitoring_note(records, lang), start=1):
            narrative.append([f'{self._label("verification_note", lang)} {index}', paragraph])
        for index, paragraph in enumerate(self.ai_writer.recommended_next_steps(records, lang), start=1):
            narrative.append([f'{self._label("next_steps", lang)} {index}', paragraph])

        dashboard = wb.create_sheet('Tableau de bord' if lang == 'fr' else 'Dashboard')
        dashboard.append([self._label('dimension', lang), self._label('value', lang)])
        self._header_style(dashboard[1], fill=ORANGE)
        for label, value in self._dashboard_rows(records, lang):
            dashboard.append([label, value])

        activity_sheet = wb.create_sheet('Synthèse activité' if lang == 'fr' else 'Activity synthesis')
        activity_sheet.freeze_panes = 'A2'
        activity_sheet.append([
            self._label('activity', lang),
            self._label('evidence_items', lang),
            self._label('supporting_media', lang),
            self._label('outputs_covered', lang),
            self._label('locations', lang),
            self._label('narrative_summary', lang),
        ])
        self._header_style(activity_sheet[1], fill=BLUE)
        grouped_by_activity: dict[str, list[ItemRecord]] = defaultdict(list)
        for record in records:
            if record.kind != 'evidence':
                continue
            grouped_by_activity[_safe_text(record.raw.get('activity'), self._label('unspecified_activity', lang))].append(record)
        for activity, grouped_records in sorted(grouped_by_activity.items(), key=lambda item: (-len(item[1]), item[0])):
            outputs = sorted({value for value in (_safe_text(r.raw.get('output')) for r in grouped_records) if value})
            locations = sorted({value for value in (_safe_text(r.raw.get('locationLabel')) for r in grouped_records) if value})
            activity_sheet.append([
                activity,
                len(grouped_records),
                sum(len(r.media_files) for r in grouped_records),
                _comma_join(outputs),
                _comma_join(locations),
                self.ai_writer.implementation_summary(activity, grouped_records, lang),
            ])

        timeline = wb.create_sheet('Chronologie' if lang == 'fr' else 'Timeline')
        timeline.freeze_panes = 'A2'
        timeline.append([
            self._label('date', lang),
            self._label('kind', lang),
            self._label('subtype', lang),
            self._label('title', lang),
            self._label('activity', lang),
            self._label('output_or_quote', lang),
            self._label('location_or_beneficiary', lang),
            self._label('media', lang),
            self._label('folder', lang),
        ])
        self._header_style(timeline[1], fill=BLUE)
        for record in records:
            timeline.append([
                record.created_at,
                self._kind_label(record.kind, lang),
                self.ai_writer._subtype_phrase(record, lang),
                record.title,
                _safe_text(record.raw.get('activity')),
                _safe_text(record.raw.get('output') if record.kind == 'evidence' else record.raw.get('quote')),
                record.location,
                len(record.media_files),
                record.relative_folder,
            ])

        evidence_sheet = wb.create_sheet('Preuves' if lang == 'fr' else 'Evidence')
        evidence_sheet.freeze_panes = 'A2'
        evidence_sheet.append([
            self._label('title', lang), self._label('created', lang), self._label('type', lang), self._label('activity', lang),
            self._label('output', lang), self._label('location', lang), self._label('description_summary', lang),
            self._label('media', lang), self._label('gps', lang), self._label('folder', lang)
        ])
        self._header_style(evidence_sheet[1], fill=BLUE)

        story_sheet = wb.create_sheet('Stories' if lang == 'fr' else 'Stories')
        story_sheet.freeze_panes = 'A2'
        story_sheet.append([
            self._label('title', lang), self._label('created', lang), self._label('beneficiary', lang), self._label('summary', lang),
            self._label('quote', lang), self._label('consent', lang), self._label('media', lang), self._label('folder', lang)
        ])
        self._header_style(story_sheet[1], fill=ORANGE)

        media_sheet = wb.create_sheet('Médias' if lang == 'fr' else 'Media')
        media_sheet.freeze_panes = 'A2'
        media_sheet.append([
            self._label('linked_item', lang), self._label('kind', lang), self._label('file_name', lang),
            self._label('type', lang), self._label('caption', lang), self._label('folder', lang)
        ])
        self._header_style(media_sheet[1], fill=BLUE)

        for record in records:
            if record.kind == 'evidence':
                evidence_sheet.append([
                    record.title,
                    record.created_at,
                    self.ai_writer._subtype_phrase(record, lang),
                    _safe_text(record.raw.get('activity')),
                    _safe_text(record.raw.get('output')),
                    record.location,
                    record.description,
                    len(record.media_files),
                    self._gps_value(record),
                    record.relative_folder,
                ])
            else:
                story_sheet.append([
                    record.title,
                    record.created_at,
                    record.location,
                    record.description,
                    _safe_text(record.raw.get('quote')),
                    self._yes_no(bool(record.raw.get('consentGiven')), lang),
                    len(record.media_files),
                    record.relative_folder,
                ])
            for media_file in record.media_files:
                media_sheet.append([
                    record.title,
                    self._kind_label(record.kind, lang),
                    media_file.name,
                    self._file_type_label(media_file, lang),
                    self.ai_writer.media_caption(record, media_file, lang),
                    record.relative_folder,
                ])

        for sheet in (overview, narrative, dashboard, activity_sheet, timeline, evidence_sheet, story_sheet, media_sheet):
            self._autosize(sheet)
            self._apply_sheet_polish(sheet, auto_filter=True)

        wb.save(output)

    def _build_project_excel_report(self, output: Path, project_code: str, project_name: str, records: list[ItemRecord], lang: str) -> None:
        self.premium_excel_report_builder.build(output, project_code, project_name, records, lang)
        return

    def _build_project_pptx_report(self, output: Path, project_code: str, project_name: str, records: list[ItemRecord], lang: str) -> None:
        self.premium_pptx_report_builder.build(output, project_code, project_name, records, lang)
        return

    def _build_project_report(self, output: Path, project_code: str, project_name: str, records: list[ItemRecord], lang: str) -> None:
        self.premium_report_builder.build(output, project_code, project_name, records, lang)
        return

        document = Document()
        self._configure_document(document)
        meta = self._project_meta(records, preferred_language=lang)
        evidence_records = [record for record in records if record.kind == 'evidence']
        story_records = [record for record in records if record.kind == 'story']

        self._cover_block(document, project_code, project_name, meta, records, lang)

        document.add_page_break()
        self._add_section_title(document, self._label('executive_summary', lang))
        for paragraph in self.ai_writer.executive_summary(project_name, meta['donor_name'], meta['country'], records, lang):
            self._body_paragraph(document, paragraph)

        self._add_section_title(document, self._label('portfolio_snapshot', lang))
        summary_table = document.add_table(rows=0, cols=2)
        self._set_table_style(summary_table)
        for label, value in [
            (self._label('donor', lang), meta['donor_name'] or self._label('not_specified', lang)),
            (self._label('country', lang), meta['country'] or self._label('not_specified', lang)),
            (self._label('reporting_window', lang), meta['reporting_window']),
            (self._label('report_language', lang), self._label('language_french', lang) if lang == 'fr' else self._label('language_english', lang)),
            (self._label('evidence_items', lang), str(len(evidence_records))),
            (self._label('story_items', lang), str(len(story_records))),
            (self._label('media_assets', lang), str(sum(len(r.media_files) for r in records))),
            (self._label('top_activity', lang), self._top_value([_safe_text(r.raw.get('activity')) for r in evidence_records], self._label('not_specified', lang))),
            (self._label('top_output', lang), self._top_value([_safe_text(r.raw.get('output')) for r in evidence_records], self._label('not_specified', lang))),
        ]:
            row = summary_table.add_row().cells
            row[0].text = label
            row[1].text = value
        self._shade_first_column(summary_table, SOFT_BLUE)

        self._add_section_title(document, self._label('key_highlights', lang))
        for bullet in self.ai_writer.key_highlights(records, lang):
            self._bullet_paragraph(document, bullet)

        self._add_section_title(document, self._label('implementation_by_activity', lang))
        activity_groups: dict[str, list[ItemRecord]] = defaultdict(list)
        for record in evidence_records:
            activity_groups[_safe_text(record.raw.get('activity'), self._label('unspecified_activity', lang))].append(record)
        if not activity_groups:
            self._body_paragraph(document, self._label('no_evidence_synced', lang))
        for activity, grouped_records in sorted(activity_groups.items(), key=lambda item: (-len(item[1]), item[0])):
            self._subheading(document, activity)
            self._body_paragraph(document, self.ai_writer.implementation_summary(activity, grouped_records, lang))
            outputs = sorted({value for value in (_safe_text(r.raw.get('output')) for r in grouped_records) if value})
            locations = sorted({value for value in (_safe_text(r.raw.get('locationLabel')) for r in grouped_records) if value})
            detail_table = document.add_table(rows=0, cols=2)
            self._set_table_style(detail_table)
            for label, value in [
                (self._label('evidence_items', lang), str(len(grouped_records))),
                (self._label('supporting_media', lang), str(sum(len(r.media_files) for r in grouped_records))),
                (self._label('outputs_covered', lang), _comma_join(outputs)),
                (self._label('locations', lang), _comma_join(locations)),
            ]:
                row = detail_table.add_row().cells
                row[0].text = label
                row[1].text = value
            self._shade_first_column(detail_table, SOFT_ORANGE)
            representative = [record.title for record in grouped_records[:4]]
            if representative:
                self._body_paragraph(
                    document,
                    _sentence(
                        (f'Éléments représentatifs : {", ".join(representative)}.' if lang == 'fr' else f'Representative entries: {", ".join(representative)}.')
                    ),
                )
            document.add_paragraph('')

        self._add_section_title(document, self._label('strategic_narrative', lang))
        for activity, grouped_records in sorted(activity_groups.items(), key=lambda item: (-len(item[1]), item[0]))[:3]:
            self._subheading(document, activity)
            for paragraph in self.ai_writer.activity_deep_dive(activity, grouped_records, lang):
                self._body_paragraph(document, paragraph)

        self._add_section_title(document, self._label('verification_note', lang))
        for paragraph in self.ai_writer.monitoring_note(records, lang):
            self._body_paragraph(document, paragraph)

        self._add_section_title(document, self._label('next_steps', lang))
        for paragraph in self.ai_writer.recommended_next_steps(records, lang):
            self._bullet_paragraph(document, paragraph)

        self._add_section_title(document, self._label('evidence_highlights', lang))
        if not evidence_records:
            self._body_paragraph(document, self._label('no_evidence_available', lang))
        for index, record in enumerate(self._prioritized_records(evidence_records)[:8], start=1):
            self._record_showcase(document, record, index=index, lang=lang)

        self._add_section_title(document, self._label('beneficiary_stories', lang))
        if not story_records:
            self._body_paragraph(document, self._label('no_stories_available', lang))
        for index, record in enumerate(self._prioritized_records(story_records)[:6], start=1):
            self._record_showcase(document, record, index=index, lang=lang)

        self._add_section_title(document, self._label('visual_annex', lang))
        image_entries = [(record, media_file) for record in self._prioritized_records(records) for media_file in record.media_files if media_file.suffix.lower() in IMAGE_EXTENSIONS]
        if image_entries:
            self._add_image_gallery(document, image_entries, lang)
        else:
            self._body_paragraph(document, self._label('no_media_annex', lang))

        video_entries = [(record, media_file) for record in self._prioritized_records(records) for media_file in record.media_files if media_file.suffix.lower() in VIDEO_EXTENSIONS]
        if video_entries:
            self._subheading(document, self._label('linked_videos', lang))
            video_table = document.add_table(rows=1, cols=3)
            self._set_table_style(video_table)
            headers = [self._label('linked_item', lang), self._label('file_name', lang), self._label('caption', lang)]
            for cell, header in zip(video_table.rows[0].cells, headers):
                cell.text = header
            self._shade_row(video_table.rows[0], BLUE, 'FFFFFF')
            for record, media_file in video_entries:
                row = video_table.add_row().cells
                row[0].text = record.title
                row[1].text = media_file.name
                row[2].text = self.ai_writer.media_caption(record, media_file, lang)

        document.save(output)

    def _build_item_doc(self, record: ItemRecord, filename: str, lang: str) -> None:
        output = record.metadata_path.parent / filename
        document = Document()
        self._configure_document(document)
        self._cover_title(document, record.title)
        self._subtitle_line(document, f'{record.project_name} ({record.project_code})')
        self._subtitle_line(document, f'{self._label("report_language", lang)}: {self._label("language_french", lang) if lang == "fr" else self._label("language_english", lang)}')
        self._meta_strip(document, [
            (self._label('created', lang), record.created_at or '-'),
            (self._label('kind', lang), self._kind_label(record.kind, lang)),
            (self._label('type', lang), self.ai_writer._subtype_phrase(record, lang)),
            (self._label('linked_folder', lang), record.relative_folder),
        ])

        self._add_section_title(document, self._label('donor_facing_narrative', lang))
        narrative = self.ai_writer.evidence_narrative(record, lang) if record.kind == 'evidence' else self.ai_writer.story_narrative(record, lang)
        self._body_paragraph(document, narrative)

        self._add_section_title(document, self._label('key_takeaways', lang))
        for item in self.ai_writer.item_takeaways(record, lang):
            self._bullet_paragraph(document, item)

        self._add_section_title(document, self._label('source_details', lang))
        source_table = document.add_table(rows=0, cols=2)
        self._set_table_style(source_table)
        details = [
            (self._label('project', lang), f'{record.project_name} ({record.project_code})'),
            (self._label('created', lang), record.created_at or '-'),
            (self._label('description_summary', lang), record.description or '-'),
        ]
        if record.kind == 'evidence':
            details.extend([
                (self._label('activity', lang), _safe_text(record.raw.get('activity'), '-')),
                (self._label('output', lang), _safe_text(record.raw.get('output'), '-')),
                (self._label('location', lang), _safe_text(record.raw.get('locationLabel'), '-')),
                (self._label('gps', lang), self._gps_value(record)),
            ])
        else:
            details.extend([
                (self._label('beneficiary', lang), _safe_text(record.raw.get('beneficiaryAlias'), '-')),
                (self._label('consent', lang), self._yes_no(bool(record.raw.get('consentGiven')), lang)),
                (self._label('quote', lang), _safe_text(record.raw.get('quote'), '-')),
            ])
        for label, value in details:
            row = source_table.add_row().cells
            row[0].text = label
            row[1].text = value
        self._shade_first_column(source_table, SOFT_BLUE)

        image_entries = [(record, path) for path in record.media_files if path.suffix.lower() in IMAGE_EXTENSIONS]
        video_files = [path for path in record.media_files if path.suffix.lower() in VIDEO_EXTENSIONS]

        if image_entries:
            self._add_section_title(document, self._label('illustrative_images', lang))
            self._add_image_gallery(document, image_entries, lang)

        if video_files:
            self._add_section_title(document, self._label('linked_videos', lang))
            for video_path in video_files:
                self._bullet_paragraph(document, self.ai_writer.media_caption(record, video_path, lang))

        document.save(output)

    def _add_image_gallery(self, document: Document, image_entries: list[tuple[ItemRecord, Path]], lang: str) -> None:
        gallery_chunks = [image_entries[index:index + 4] for index in range(0, len(image_entries), 4)]
        for chunk_index, chunk in enumerate(gallery_chunks):
            if chunk_index > 0:
                document.add_page_break()
                self._subheading(document, self._label('visual_annex_continued', lang))
            table = document.add_table(rows=2, cols=2)
            self._set_table_style(table)
            flat_cells = [cell for row in table.rows for cell in row.cells]
            for cell in flat_cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                self._shade_cell(cell, SOFT_GREY)
            for cell, entry in zip(flat_cells, chunk):
                record, image_path = entry
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                try:
                    run.add_picture(str(image_path), width=Inches(1.95))
                except Exception:
                    paragraph.add_run(image_path.name)
                caption = cell.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = caption.add_run(self.ai_writer.media_caption(record, image_path, lang))
                cap_run.italic = True
                cap_run.font.size = Pt(8.5)
                cap_run.font.color.rgb = RGBColor.from_string(MID_GREY)
            for cell in flat_cells[len(chunk):]:
                cell.text = ''
            document.add_paragraph('')

    def _record_showcase(self, document: Document, record: ItemRecord, index: int, lang: str) -> None:
        label = self._label('evidence_item', lang) if record.kind == 'evidence' else self._label('story_item', lang)
        self._subheading(document, f'{label} {index:02d} — {record.title}')
        meta_table = document.add_table(rows=0, cols=2)
        self._set_table_style(meta_table)
        rows = [
            (self._label('created', lang), record.created_at or '-'),
            (self._label('location_or_beneficiary', lang), record.location or '-'),
            (self._label('media_assets', lang), str(len(record.media_files))),
        ]
        if record.kind == 'evidence':
            rows.extend([
                (self._label('activity', lang), _safe_text(record.raw.get('activity'), '-')),
                (self._label('output', lang), _safe_text(record.raw.get('output'), '-')),
            ])
        else:
            rows.append((self._label('consent', lang), self._yes_no(bool(record.raw.get('consentGiven')), lang)))
        for label_text, value in rows:
            row = meta_table.add_row().cells
            row[0].text = label_text
            row[1].text = value
        self._shade_first_column(meta_table, SOFT_GREY)
        self._body_paragraph(document, self.ai_writer.evidence_narrative(record, lang) if record.kind == 'evidence' else self.ai_writer.story_narrative(record, lang))
        for bullet in self.ai_writer.item_takeaways(record, lang):
            self._bullet_paragraph(document, bullet)
        image_files = [path for path in record.media_files if path.suffix.lower() in IMAGE_EXTENSIONS]
        if image_files:
            gallery_table = document.add_table(rows=1, cols=min(2, len(image_files[:2])))
            self._set_table_style(gallery_table)
            for cell, image_path in zip(gallery_table.rows[0].cells, image_files[:2]):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    paragraph.add_run().add_picture(str(image_path), width=Inches(1.95))
                except Exception:
                    paragraph.add_run(image_path.name)
        quote = _safe_text(record.raw.get('quote'))
        if quote:
            quote_para = document.add_paragraph()
            quote_para.style = document.styles['Quote']
            quote_para.add_run(f'« {quote} »' if lang == 'fr' else f'“{quote}”')
        document.add_paragraph('')

    def _project_meta(self, records: list[ItemRecord], preferred_language: str | None = None) -> dict[str, str]:
        project = {}
        for record in records:
            candidate = record.raw.get('project')
            if isinstance(candidate, dict) and candidate:
                project = candidate
                break
        start_values = [dt for dt in (_parse_iso(r.raw.get('createdAt')) for r in records) if dt is not None]
        detected_lang = self._detect_project_language(records, preferred_language)
        reporting_window = '-'
        if start_values:
            connector = 'au' if detected_lang == 'fr' else 'to'
            reporting_window = f'{min(start_values).strftime("%Y-%m-%d")} {connector} {max(start_values).strftime("%Y-%m-%d")}'
        polished_project = self._polish_project_payload(project, detected_lang)
        return {
            'donor_name': _safe_text(polished_project.get('donorName')),
            'country': _safe_text(polished_project.get('country')),
            'reporting_window': reporting_window,
            'language': detected_lang,
        }

    def _dashboard_rows(self, records: list[ItemRecord], lang: str) -> list[tuple[str, str]]:
        evidence_records = [record for record in records if record.kind == 'evidence']
        story_records = [record for record in records if record.kind == 'story']
        activity_counts = Counter(_safe_text(record.raw.get('activity')) for record in evidence_records if _safe_text(record.raw.get('activity')))
        output_counts = Counter(_safe_text(record.raw.get('output')) for record in evidence_records if _safe_text(record.raw.get('output')))
        location_counts = Counter(record.location for record in records if record.location and record.location != '-')
        rows = [
            (self._label('evidence_items', lang), str(len(evidence_records))),
            (self._label('story_items', lang), str(len(story_records))),
            (self._label('media_assets', lang), str(sum(len(record.media_files) for record in records))),
            (self._label('gps_enabled_evidence', lang), str(sum(1 for record in evidence_records if record.raw.get('latitude') is not None and record.raw.get('longitude') is not None))),
        ]
        prefix_activity = 'Activité · ' if lang == 'fr' else 'Activity · '
        prefix_output = 'Output · '
        prefix_location = 'Lieu · ' if lang == 'fr' else 'Location · '
        for activity, count in activity_counts.most_common(5):
            rows.append((f'{prefix_activity}{activity}', str(count)))
        for output, count in output_counts.most_common(5):
            rows.append((f'{prefix_output}{output}', str(count)))
        for location, count in location_counts.most_common(5):
            rows.append((f'{prefix_location}{location}', str(count)))
        return rows

    def _prioritized_records(self, records: list[ItemRecord]) -> list[ItemRecord]:
        def sort_key(record: ItemRecord):
            parsed = _parse_iso(record.raw.get('createdAt'))
            timestamp = parsed.timestamp() if parsed else 0
            return (len(record.media_files), timestamp, len(_safe_text(record.description)), record.title.lower())
        return sorted(records, key=sort_key, reverse=True)

    def _record_language(self, record: ItemRecord, fallback: str = 'fr') -> str:
        return _normalize_lang(record.raw.get('reportLanguage') or record.raw.get('languageCode') or fallback)

    def _detect_project_language(self, records: list[ItemRecord], preferred_language: str | None = None) -> str:
        if preferred_language:
            return _normalize_lang(preferred_language)
        dated: list[tuple[float, str]] = []
        for record in records:
            lang = self._record_language(record, fallback='fr')
            moment = _parse_iso(record.raw.get('syncedAt')) or _parse_iso(record.raw.get('createdAt'))
            timestamp = moment.timestamp() if moment else 0
            dated.append((timestamp, lang))
        if not dated:
            return 'fr'
        dated.sort(reverse=True)
        return dated[0][1]

    def _gps_value(self, record: ItemRecord) -> str:
        lat = record.raw.get('latitude')
        lon = record.raw.get('longitude')
        if lat is None or lon is None:
            return '-'
        return f'{lat}, {lon}'

    def _top_value(self, values: list[str], fallback: str = 'Not specified') -> str:
        clean = [value for value in values if value]
        if not clean:
            return fallback
        return Counter(clean).most_common(1)[0][0]

    def _configure_document(self, document: Document) -> None:
        section = document.sections[0]
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.72)
        section.right_margin = Inches(0.72)
        styles = document.styles
        styles['Normal'].font.name = 'Aptos'
        styles['Normal'].font.size = Pt(10.5)
        styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Aptos')
        styles['Heading 1'].font.name = 'Aptos Display'
        styles['Heading 1'].font.size = Pt(18)
        styles['Heading 1'].font.bold = True
        styles['Heading 1'].font.color.rgb = RGBColor.from_string(BLUE)
        styles['Heading 1'].paragraph_format.space_before = Pt(10)
        styles['Heading 1'].paragraph_format.space_after = Pt(8)
        styles['Heading 2'].font.name = 'Aptos'
        styles['Heading 2'].font.size = Pt(13)
        styles['Heading 2'].font.bold = True
        styles['Heading 2'].font.color.rgb = RGBColor.from_string(BLUE)
        styles['Heading 2'].paragraph_format.space_before = Pt(8)
        styles['Heading 2'].paragraph_format.space_after = Pt(4)
        styles['Quote'].font.name = 'Aptos'
        styles['Quote'].font.italic = True
        styles['Quote'].font.color.rgb = RGBColor.from_string(DARK)

    def _cover_block(self, document: Document, project_code: str, project_name: str, meta: dict[str, str], records: list[ItemRecord], lang: str) -> None:
        kicker = document.add_paragraph()
        kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kicker_run = kicker.add_run(self._label('cover_kicker', lang))
        kicker_run.font.name = 'Aptos'
        kicker_run.font.size = Pt(9.5)
        kicker_run.bold = True
        kicker_run.font.color.rgb = RGBColor.from_string(ORANGE)

        self._cover_title(document, self._label('cover_title', lang))
        self._subtitle_line(document, project_name)
        self._subtitle_line(document, f'{self._label("project_code", lang)}: {project_code}')
        self._subtitle_line(document, f'{self._label("donor", lang)}: {meta["donor_name"] or self._label("not_specified", lang)} • {self._label("country", lang)}: {meta["country"] or self._label("not_specified", lang)}')
        self._subtitle_line(document, f'{self._label("reporting_window", lang)}: {meta["reporting_window"]}')
        self._subtitle_line(document, f'{self._label("report_language", lang)}: {self._label("language_french", lang) if lang == "fr" else self._label("language_english", lang)}')
        self._subtitle_line(document, f'{self._label("generated_on", lang)} {datetime.now().strftime("%Y-%m-%d %H:%M")}.')
        document.add_paragraph('')

        hero_table = document.add_table(rows=2, cols=3)
        self._set_table_style(hero_table)
        hero_values = [
            (self._label('evidence_items', lang), str(sum(1 for r in records if r.kind == 'evidence'))),
            (self._label('story_items', lang), str(sum(1 for r in records if r.kind == 'story'))),
            (self._label('media_assets', lang), str(sum(len(r.media_files) for r in records))),
            (self._label('top_activity', lang), self._top_value([_safe_text(r.raw.get('activity')) for r in records if r.kind == 'evidence'], self._label('not_specified', lang))),
            (self._label('top_output', lang), self._top_value([_safe_text(r.raw.get('output')) for r in records if r.kind == 'evidence'], self._label('not_specified', lang))),
            (self._label('local_archive', lang), 'GrantProof / projects / reports'),
        ]
        for index, (label, value) in enumerate(hero_values):
            cell = hero_table.cell(index // 3, index % 3)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label_run = paragraph.add_run(label + "\n")
            label_run.bold = True
            label_run.font.name = 'Aptos'
            label_run.font.size = Pt(9.5)
            label_run.font.color.rgb = RGBColor.from_string(BLUE)
            value_run = paragraph.add_run(value)
            value_run.font.name = 'Aptos'
            value_run.font.size = Pt(11)
            value_run.font.color.rgb = RGBColor.from_string(DARK)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for row in hero_table.rows:
            for cell in row.cells:
                self._shade_cell(cell, SOFT_BLUE)
        document.add_paragraph('')

    def _cover_title(self, document: Document, text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = 'Aptos'
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor.from_string(BLUE)

    def _subtitle_line(self, document: Document, text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        run.font.name = 'Aptos'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(DARK)

    def _meta_strip(self, document: Document, pairs: list[tuple[str, str]]) -> None:
        table = document.add_table(rows=0, cols=2)
        self._set_table_style(table)
        for label, value in pairs:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value
        self._shade_first_column(table, SOFT_BLUE)
        document.add_paragraph('')

    def _add_section_title(self, document: Document, text: str) -> None:
        document.add_heading(text, level=1)

    def _subheading(self, document: Document, text: str) -> None:
        document.add_heading(text, level=2)

    def _body_paragraph(self, document: Document, text: str) -> None:
        paragraph = document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.2

    def _bullet_paragraph(self, document: Document, text: str) -> None:
        paragraph = document.add_paragraph(style='List Bullet')
        paragraph.add_run(text)
        paragraph.paragraph_format.space_after = Pt(4)

    def _set_table_style(self, table) -> None:
        table.style = 'Table Grid'
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _shade_first_column(self, table, fill_hex: str) -> None:
        for row in table.rows:
            self._shade_cell(row.cells[0], fill_hex)
            self._bold_paragraphs(row.cells[0])

    def _shade_row(self, row, fill_hex: str, text_hex: str) -> None:
        for cell in row.cells:
            self._shade_cell(cell, fill_hex)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor.from_string(text_hex)
                    run.bold = True

    def _shade_cell(self, cell, fill_hex: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill_hex)
        tc_pr.append(shd)

    def _bold_paragraphs(self, cell) -> None:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(BLUE)

    def _header_style(self, cells: Iterable, fill: str = BLUE) -> None:
        for cell in cells:
            cell.font = Font(name='Aptos', bold=True, color='FFFFFF')
            cell.fill = PatternFill('solid', fgColor=fill)
            cell.alignment = Alignment(vertical='top', wrap_text=True)
            cell.border = THIN_GREY_BORDER

    def _autosize(self, worksheet) -> None:
        for column_cells in worksheet.columns:
            max_length = 0
            for cell in column_cells:
                value = _safe_text(cell.value)
                for line in value.splitlines() or ['']:
                    max_length = max(max_length, len(line))
            worksheet.column_dimensions[get_column_letter(column_cells[0].column)].width = min(max(max_length + 2, 14), 42)

    def _apply_sheet_polish(self, worksheet, auto_filter: bool = False) -> None:
        for row_idx, row in enumerate(worksheet.iter_rows(), start=1):
            for cell in row:
                if row_idx != 1:
                    cell.font = Font(name='Aptos', size=10)
                cell.alignment = Alignment(vertical='top', wrap_text=True)
                cell.border = THIN_GREY_BORDER
            if row_idx > 1 and row_idx % 2 == 0:
                for cell in row:
                    if cell.fill.fgColor.rgb in (None, '00000000'):
                        cell.fill = PatternFill('solid', fgColor='FBFCFF')
        if auto_filter and worksheet.max_row >= 1 and worksheet.max_column >= 1:
            worksheet.auto_filter.ref = worksheet.dimensions
        worksheet.sheet_view.showGridLines = False
        worksheet.freeze_panes = worksheet.freeze_panes or 'A2'

    def _label(self, key: str, lang: str) -> str:
        labels = {
            'fr': {
                'metric': 'Indicateur', 'value': 'Valeur', 'project_code': 'Code projet', 'project_name': 'Nom du projet',
                'donor': 'Bailleur', 'country': 'Pays', 'reporting_window': 'Période couverte', 'generated_at': 'Généré le',
                'generated_on': 'Généré le', 'evidence_items': 'Preuves', 'story_items': 'Stories', 'media_assets': 'Médias', 'gps_enabled_evidence': 'Preuves GPS',
                'top_activity': 'Activité dominante', 'top_output': 'Output dominant', 'report_language': 'Langue du rapport',
                'language_french': 'Français', 'language_english': 'Anglais', 'section': 'Section', 'content': 'Contenu',
                'executive_summary': 'Résumé exécutif', 'key_highlights': 'Points clés pour le bailleur', 'verification_note': 'Note de vérification', 'strategic_narrative': 'Lecture stratégique par activité', 'next_steps': 'Pistes de renforcement',
                'dimension': 'Dimension', 'date': 'Date', 'kind': 'Nature', 'subtype': 'Sous-type', 'title': 'Titre', 'activity': 'Activité',
                'output_or_quote': 'Output / citation', 'location_or_beneficiary': 'Lieu / bénéficiaire', 'media': 'Médias', 'folder': 'Dossier',
                'created': 'Créé le', 'type': 'Type', 'output': 'Output', 'location': 'Lieu', 'description_summary': 'Description / résumé',
                'gps': 'GPS', 'beneficiary': 'Bénéficiaire', 'summary': 'Résumé', 'quote': 'Citation', 'consent': 'Consentement',
                'linked_item': 'Élément lié', 'file_name': 'Nom du fichier', 'caption': 'Légende', 'portfolio_snapshot': 'Vue d’ensemble du portefeuille',
                'implementation_by_activity': 'Narratif de mise en œuvre par activité', 'supporting_media': 'Médias de support', 'outputs_covered': 'Outputs couverts',
                'locations': 'Lieux', 'no_evidence_synced': 'Aucune preuve n’a encore été synchronisée pour ce projet.',
                'evidence_highlights': 'Preuves mises en avant', 'beneficiary_stories': 'Stories et signaux qualitatifs',
                'no_evidence_available': 'Aucune preuve n’est disponible pour le moment.', 'no_stories_available': 'Aucune story n’est disponible pour le moment.',
                'visual_annex': 'Annexe visuelle', 'visual_annex_continued': 'Annexe visuelle (suite)', 'no_media_annex': 'Aucune annexe média n’est disponible pour cette période.',
                'linked_videos': 'Vidéos liées', 'donor_facing_narrative': 'Narratif bailleur', 'key_takeaways': 'Points clés', 'source_details': 'Détails source',
                'project': 'Projet', 'illustrative_images': 'Galerie visuelle', 'cover_title': 'GrantProof — Rapport projet premium', 'cover_kicker': 'Pack narratif terrain • version premium', 'narrative_summary': 'Synthèse narrative', 'not_specified': 'Non renseigné',
                'local_archive': 'Archive locale', 'linked_folder': 'Dossier lié', 'evidence_item': 'Preuve', 'story_item': 'Story', 'unspecified_activity': 'Activité non précisée'
            },
            'en': {
                'metric': 'Metric', 'value': 'Value', 'project_code': 'Project code', 'project_name': 'Project name',
                'donor': 'Donor', 'country': 'Country', 'reporting_window': 'Reporting window', 'generated_at': 'Generated at',
                'generated_on': 'Generated on', 'evidence_items': 'Evidence items', 'story_items': 'Story items', 'media_assets': 'Media assets', 'gps_enabled_evidence': 'GPS-enabled evidence',
                'top_activity': 'Top activity', 'top_output': 'Top output', 'report_language': 'Report language',
                'language_french': 'French', 'language_english': 'English', 'section': 'Section', 'content': 'Content',
                'executive_summary': 'Executive summary', 'key_highlights': 'Key highlights for donor reporting', 'verification_note': 'Verification note', 'strategic_narrative': 'Strategic reading by activity', 'next_steps': 'Recommended next steps',
                'dimension': 'Dimension', 'date': 'Date', 'kind': 'Kind', 'subtype': 'Subtype', 'title': 'Title', 'activity': 'Activity',
                'output_or_quote': 'Output / quote', 'location_or_beneficiary': 'Location / beneficiary', 'media': 'Media', 'folder': 'Folder',
                'created': 'Created', 'type': 'Type', 'output': 'Output', 'location': 'Location', 'description_summary': 'Description / summary',
                'gps': 'GPS', 'beneficiary': 'Beneficiary', 'summary': 'Summary', 'quote': 'Quote', 'consent': 'Consent',
                'linked_item': 'Linked item', 'file_name': 'File name', 'caption': 'Caption', 'portfolio_snapshot': 'Portfolio snapshot',
                'implementation_by_activity': 'Implementation narrative by activity', 'supporting_media': 'Supporting media', 'outputs_covered': 'Outputs covered',
                'locations': 'Locations', 'no_evidence_synced': 'No evidence has been synchronized yet for this project.',
                'evidence_highlights': 'Evidence highlights', 'beneficiary_stories': 'Beneficiary stories and qualitative signals',
                'no_evidence_available': 'No evidence items are available yet.', 'no_stories_available': 'No beneficiary stories are available yet.',
                'visual_annex': 'Visual annex', 'visual_annex_continued': 'Visual annex (continued)', 'no_media_annex': 'No media annex is available for this reporting window.',
                'linked_videos': 'Linked videos', 'donor_facing_narrative': 'Donor-facing narrative', 'key_takeaways': 'Key takeaways', 'source_details': 'Source details',
                'project': 'Project', 'illustrative_images': 'Visual gallery', 'cover_title': 'GrantProof — Premium project report', 'cover_kicker': 'Field evidence pack • premium edition', 'narrative_summary': 'Narrative summary', 'not_specified': 'Not specified',
                'local_archive': 'Local archive', 'linked_folder': 'Linked folder', 'evidence_item': 'Evidence', 'story_item': 'Story', 'unspecified_activity': 'Unspecified activity'
            },
        }
        return labels[_normalize_lang(lang)][key]

    def _kind_label(self, kind: str, lang: str) -> str:
        if lang == 'fr':
            return 'Preuve' if kind == 'evidence' else 'Story'
        return 'Evidence' if kind == 'evidence' else 'Story'

    def _file_type_label(self, file_path: Path, lang: str) -> str:
        if file_path.suffix.lower() in IMAGE_EXTENSIONS:
            return 'Image' if lang == 'fr' else 'Image'
        if file_path.suffix.lower() in VIDEO_EXTENSIONS:
            return 'Vidéo' if lang == 'fr' else 'Video'
        return 'Fichier' if lang == 'fr' else 'File'

    def _yes_no(self, value: bool, lang: str) -> str:
        return 'Oui' if value and lang == 'fr' else 'Yes' if value else 'Non' if lang == 'fr' else 'No'
