from __future__ import annotations

import math
import re
import tempfile
import zipfile
import sys
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any

import shapefile  # pyshp
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageChops, ImageDraw, ImageFont, ImageOps, ImageStat, ImageFilter

IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.webp'}
VIDEO_EXTENSIONS = {'.mp4', '.mov', '.m4v', '.avi', '.mkv', '.webm'}

BLUE = '064890'
BLUE_2 = '0A57B5'
BLUE_3 = '1665C1'
PALE_BLUE = 'EEF5FF'
SOFT_BLUE = 'F6FAFF'
ORANGE = 'F56D28'
DARK = '1D2D44'
MUTED = '5D6B7C'
BORDER = 'D9E4F2'
BORDER_LIGHT = 'E8EEF7'
WHITE = 'FFFFFF'
LIGHT_GREY = 'F4F6FA'

COUNTRY_ALIASES = {
    'niger': 'Niger',
    'france': 'France',
    'république démocratique du congo': 'Dem. Rep. Congo',
    'republique democratique du congo': 'Dem. Rep. Congo',
    'rd congo': 'Dem. Rep. Congo',
    'rdc': 'Dem. Rep. Congo',
    'democratic republic of congo': 'Dem. Rep. Congo',
    'drc': 'Dem. Rep. Congo',
    'burkina faso': 'Burkina Faso',
    'mali': 'Mali',
    'tchad': 'Chad',
    'chad': 'Chad',
    'nigeria': 'Nigeria',
    'sénégal': 'Senegal',
    'senegal': 'Senegal',
    'cameroun': 'Cameroon',
    'cameroon': 'Cameroon',
    'côte d’ivoire': "Côte d'Ivoire",
    "côte d'ivoire": "Côte d'Ivoire",
    'cote d ivoire': "Côte d'Ivoire",
    'south sudan': 'S. Sudan',
    'soudan du sud': 'S. Sudan',
    'sudan': 'Sudan',
}

COUNTRY_DISPLAY_FR = {
    'Niger': 'Niger',
    'France': 'France',
    'Dem. Rep. Congo': 'RD Congo',
    'Burkina Faso': 'Burkina Faso',
    'Mali': 'Mali',
    'Chad': 'Tchad',
    'Nigeria': 'Nigeria',
    'Senegal': 'Sénégal',
    'Cameroon': 'Cameroun',
    "Côte d'Ivoire": 'Côte d’Ivoire',
    'S. Sudan': 'Soudan du Sud',
    'Sudan': 'Soudan',
}

HUMANITARIAN_ICON_MAP = {
    # Legacy compatibility: all entries now point to the supplied V3 library filenames.
    'food': 'food_security.png',
    'food_security': 'food_security.png',
    'agriculture': 'agriculture.png',
    'shelter': 'shelter.png',
    'wash': 'wash.png',
    'health': 'health.png',
    'nutrition': 'nutrition.png',
    'education': 'education.png',
    'protection': 'protection.png',
    'cash': 'cash_assistance.png',
    'coordination': 'cccm.png',
    'people': 'reunion_communautaire.png',
    'beneficiary': 'reunion_communautaire.png',
    'beneficiaries': 'reunion_communautaire.png',
    'group': 'reunion_communautaire.png',
    'groups': 'reunion_communautaire.png',
    'kits': 'distribution_semences.png',
    'box': 'entrepot_stock.png',
    'media': 'preuve_photo_camera.png',
    'photo': 'preuve_photo_camera.png',
    'camera': 'preuve_photo_camera.png',
    'check': 'checklist_conformite.png',
    'evidence': 'checklist_conformite.png',
    'gps': 'cartographie_gps.png',
    'pin': 'cartographie_gps.png',
}

REPORT_ICON_MAP = {
    # Fixed section-title icons stay as designed in the model.
    'section_overview': 'model_section_overview',
    'section_kpis': 'model_section_kpis',
    'section_logframe': 'model_section_logframe',
    'section_sector': 'model_section_sector',
    'section_highlights': 'model_section_highlights',
    'section_location': 'model_section_location',
    'model_section_overview': 'model_section_overview',
    'model_section_kpis': 'model_section_kpis',
    'model_section_sector': 'model_section_sector',
    'model_section_highlights': 'model_section_highlights',
    'model_section_location': 'model_section_location',
    'model_section_logframe': 'model_section_logframe',

    # Header meta icons are resolved to the supplied V3 library, then recolored white at runtime.
    'meta_org': 'reunion_communautaire',
    'meta_pin': 'cartographie_gps',
    'meta_calendar': 'suivi_evaluation',

    # People / target groups / protection.
    'beneficiary': 'reunion_communautaire',
    'beneficiaries': 'reunion_communautaire',
    'people': 'reunion_communautaire',
    'person': 'reunion_communautaire',
    'participant': 'formation_atelier',
    'participants': 'formation_atelier',
    'group': 'reunion_communautaire',
    'groups': 'reunion_communautaire',
    'groupements': 'reunion_communautaire',
    'household': 'shelter',
    'households': 'shelter',
    'menages': 'shelter',
    'ménages': 'shelter',
    'women': 'autonomisation_femmes',
    'children': 'protection_enfance',
    'elderly': 'soutien_personnes_agees',
    'disability': 'inclusion_handicap',
    'gbv': 'prevention_vbg',
    'psychosocial': 'soutien_psychosocial',
    'legal': 'assistance_juridique',

    # Sectors from the supplied V3 library.
    'food': 'food_security',
    'food_security': 'food_security',
    'distribution_food': 'distribution_alimentaire',
    'agriculture': 'agriculture',
    'farmer': 'agriculture',
    'farmers': 'agriculture',
    'agriculteurs': 'agriculture',
    'seeds': 'distribution_semences',
    'irrigation': 'irrigation',
    'livestock': 'elevage',
    'forestry': 'foresterie',
    'fishing': 'peche',
    'garden': 'jardin_potager',
    'market': 'marche_local',
    'wash': 'wash',
    'water': 'forage_puits',
    'well': 'forage_puits',
    'reservoir': 'reservoir_eau',
    'latrine': 'latrines',
    'hygiene': 'lavage_des_mains',
    'health': 'health',
    'clinic': 'centre_de_sante',
    'nutrition': 'nutrition',
    'education': 'education',
    'school': 'ecole',
    'protection': 'protection',
    'shelter': 'shelter',
    'construction': 'construction',
    'logistics': 'logistics',
    'coordination': 'cccm',
    'cccm': 'cccm',
    'early_recovery': 'early_recovery',
    'etc': 'etc',
    'cash': 'cash_assistance',
    'cash_transfer': 'cash_assistance',
    'cash_assistance': 'cash_assistance',
    'voucher': 'cash_assistance',

    # Evidence, monitoring, reporting, mapping and compliance.
    'kit': 'distribution_semences',
    'kits': 'distribution_semences',
    'box': 'entrepot_stock',
    'distribution': 'distribution_alimentaire',
    'media': 'preuve_photo_camera',
    'photo': 'preuve_photo_camera',
    'camera': 'preuve_photo_camera',
    'evidence': 'checklist_conformite',
    'check': 'checklist_conformite',
    'doc': 'rapport_document',
    'document': 'rapport_document',
    'report': 'rapport_document',
    'monitoring': 'suivi_evaluation',
    'evaluation': 'suivi_evaluation',
    'data': 'collecte_de_donnees',
    'survey': 'enquete_questionnaire',
    'dashboard': 'tableau_de_bord',
    'table': 'tableau_de_donnees',
    'bar': 'graphique_barres',
    'line_chart': 'graphique_lignes',
    'pie_chart': 'graphique_circulaire',
    'budget': 'budget_finances',
    'procurement': 'passation_de_marches',
    'warehouse': 'entrepot_stock',
    'cold_chain': 'chaine_du_froid',
    'fleet': 'transport_flotte',
    'gps': 'cartographie_gps',
    'pin': 'cartographie_gps',
    'location': 'cartographie_gps',
    'map': 'cartographie_gps',
    'internet': 'connectivite_internet',
    'solar': 'energie_solaire',
    'electricity': 'electricite_generateur',
    'road': 'route_pont',
    'bridge': 'route_pont',
    'environment': 'biodiversite_faune',
    'anti_poaching': 'lutte_anti_braconnage',
    'activity': 'rapport_document',
    'implementation': 'suivi_evaluation',
    'success': 'checklist_conformite',
    'training': 'formation_atelier',
    'awareness': 'sensibilisation',

    # Compatibility with model keys returned by the report logic: map them to supplied V3 icons.
    'model_kpi_beneficiaries': 'reunion_communautaire',
    'model_kpi_media': 'preuve_photo_camera',
    'model_kpi_evidence': 'checklist_conformite',
    'model_kpi_kits': 'distribution_semences',
    'model_kpi_groups': 'reunion_communautaire',
    'model_kpi_cash': 'cash_assistance',
    'model_kpi_shelter': 'shelter',
    'model_kpi_training': 'formation_atelier',
    'model_info_location': 'cartographie_gps',
    'model_info_people': 'reunion_communautaire',
    'model_info_photo': 'preuve_photo_camera',
    'model_sector_food': 'food_security',
    'model_sector_agriculture': 'agriculture',
}

ACTIVITY_ICON_KEYWORDS = [
    ('cash_assistance', ['cash', 'transfert monétaire', 'transferts monétaires', 'voucher', 'coupon', 'assistance monétaire']),
    ('distribution_semences', ['semence', 'semences', 'intrants', 'kit agricole', 'kits agricoles']),
    ('distribution_alimentaire', ['distribution alimentaire', 'vivres', 'ration', 'food distribution', 'distribution de vivres']),
    ('distribution_alimentaire', ['distribution', 'remise', 'dotation']),
    ('construction', ['construction', 'réhabilitation', 'rehabilitation', 'ouvrage', 'bâtiment', 'batiment']),
    ('shelter', ['abri', 'abris', 'shelter', 'logement', 'nfi']),
    ('formation_atelier', ['formation', 'atelier', 'session', 'renforcement de capacités', 'capacity building']),
    ('agriculture', ['agricole', 'agriculture', 'maraîcher', 'maraicher', 'agriculteur', 'agriculteurs']),
    ('irrigation', ['irrigation', 'forage', 'pompe', 'arrosage']),
    ('elevage', ['bétail', 'betail', 'élevage', 'elevage', 'livestock', 'caprin', 'ovin']),
    ('food_security', ['sécurité alimentaire', 'food security', 'moyens d’existence', 'livelihood']),
    ('wash', ['wash', 'eha', 'eau', 'latrine', 'assainissement', 'hygiène', 'hygiene']),
    ('health', ['santé', 'clinique', 'médical', 'soins', 'vaccination']),
    ('education', ['école', 'ecole', 'éducation', 'education', 'scolaire', 'enseignant', 'élève']),
    ('protection', ['protection', 'vbg', 'gbv', 'enfant', 'psychosocial']),
    ('marche_local', ['marché', 'market', 'commercialisation']),
    ('collecte_de_donnees', ['collecte de données', 'data collection', 'enquête', 'questionnaire', 'monitoring']),
    ('rapport_document', ['rapport', 'document', 'reporting', 'justificatif']),
]



INDICATOR_PATTERNS = [
    ('beneficiary', r'(\d[\d\s.,]*)\s*(?:bénéficiaires?|beneficiaires?|beneficiaries|participants?|personnes?)', 'Bénéficiaires', 'Beneficiaries', 'people', 130),
    ('kit', r'(\d[\d\s.,]*)\s*(?:kits?)', 'Kits distribués', 'Distributed kits', 'kits', 125),
    ('group', r'(\d[\d\s.,]*)\s*(?:groupements?|groupes?|groups?)', 'Groupements', 'Groups', 'groups', 120),
    ('cash', r'(\d[\d\s.,]*)\s*(?:transferts?\s+monétaires?|cash|vouchers?|coupons?)', 'Transferts', 'Cash transfers', 'cash_transfer', 118),
    ('shelter', r'(?:construction|construit|construits|réhabilitation|rehabilitation)?\s*(?:de\s*)?(\d[\d\s.,]*)\s*(?:abris?|shelters?)', 'Abris', 'Shelters', 'shelter', 125),
    ('farmer', r'(\d[\d\s.,]*)\s*(?:agriculteurs?|farmers?)', 'Agriculteurs', 'Farmers', 'farmers', 116),
    ('household', r'(\d[\d\s.,]*)\s*(?:ménages?|menages?|households?)', 'Ménages', 'Households', 'households', 112),
    ('latrine', r'(\d[\d\s.,]*)\s*(?:latrines?)', 'Latrines', 'Latrines', 'wash', 108),
    ('school', r'(\d[\d\s.,]*)\s*(?:écoles?|ecoles?|schools?)', 'Écoles', 'Schools', 'education', 104),
]

SECTOR_RULES = [
    {
        'key': 'shelter', 'fr': 'Abris', 'en': 'Shelter', 'icon': 'shelter',
        'keywords': ['abri', 'abris', 'shelter', 'logement', 'construction', 'réhabilitation', 'rehabilitation', 'maison', 'habitat', 'ame', 'nfi'],
    },
    {
        'key': 'food_security', 'fr': 'Sécurité alimentaire', 'en': 'Food security', 'icon': 'food',
        'keywords': ['sécurité alimentaire', 'food security', 'vivres', 'food', 'ration', 'maraîcher', 'maraicher', 'semence', 'intrants', 'kit agricole', 'kits agricole', 'kits agricoles', 'agricole', 'cash for food'],
    },
    {
        'key': 'agriculture', 'fr': 'Agriculture', 'en': 'Agriculture', 'icon': 'agriculture',
        'keywords': ['agriculture', 'agricultural', 'agricole', 'agricoles', 'market-gardening', 'gardening', 'maraîcher', 'maraicher', 'semence', 'irrigation', 'élevage', 'livelihood', 'moyens d’existence', 'moyens existence', 'agriculteur', 'agriculteurs'],
    },
    {
        'key': 'wash', 'fr': 'Eau, hygiène et assainissement', 'en': 'WASH', 'icon': 'wash',
        'keywords': ['wash', 'eha', 'eau', 'hygiène', 'hygiene', 'assainissement', 'latrine', 'chloration', 'forage', 'water', 'sanitation'],
    },
    {
        'key': 'health', 'fr': 'Santé', 'en': 'Health', 'icon': 'health',
        'keywords': ['santé', 'health', 'clinique', 'centre de santé', 'médicament', 'vaccination', 'choléra', 'cholera', 'soins'],
    },
    {
        'key': 'nutrition', 'fr': 'Nutrition', 'en': 'Nutrition', 'icon': 'nutrition',
        'keywords': ['nutrition', 'malnutrition', 'anjé', 'anje', 'mas', 'mam', 'dépistage', 'dépister'],
    },
    {
        'key': 'education', 'fr': 'Éducation', 'en': 'Education', 'icon': 'education',
        'keywords': ['éducation', 'education', 'école', 'ecole', 'enseignant', 'classe', 'élève', 'scolaire', 'school'],
    },
    {
        'key': 'protection', 'fr': 'Protection', 'en': 'Protection', 'icon': 'protection',
        'keywords': ['protection', 'vbg', 'gbv', 'violence', 'enfant', 'psychosocial', 'psea', 'réunification', 'protection de l’enfance'],
    },
    {
        'key': 'cash', 'fr': 'Transferts monétaires', 'en': 'Cash assistance', 'icon': 'cash',
        'keywords': ['cash', 'transfert monétaire', 'transferts monétaires', 'voucher', 'coupon', 'assistance monétaire'],
    },
    {
        'key': 'coordination', 'fr': 'Coordination', 'en': 'Coordination', 'icon': 'coordination',
        'keywords': ['coordination', 'réunion', 'meeting', 'cluster', 'atelier', 'planification'],
    },
]

QUANTITY_PATTERNS = [
    ('shelter', r'(?:construction|construit|construits|réhabilitation|rehabilitation)?\s*(?:de\s*)?(\d[\d\s.,]*)\s*(abris?|shelters?)', 'Abris', 'Shelters'),
    ('kit', r'(\d[\d\s.,]*)\s*(kits?)', 'Kits', 'Kits'),
    ('kit_context', r'(?:kit|kits)[^\d]{0,35}(\d[\d\s.,]*)', 'Kits', 'Kits'),
    ('farmer', r'(\d[\d\s.,]*)\s*(agriculteurs?|farmers?)', 'Agriculteurs', 'Farmers'),
    ('latrine', r'(\d[\d\s.,]*)\s*(latrines?)', 'Latrines', 'Latrines'),
    ('school', r'(\d[\d\s.,]*)\s*(écoles?|ecoles?|schools?)', 'Écoles', 'Schools'),
    ('household', r'(\d[\d\s.,]*)\s*(ménages?|menages?|households?)', 'Ménages', 'Households'),
    ('beneficiary', r'(\d[\d\s.,]*)\s*(bénéficiaires?|beneficiaires?|beneficiaries|participants?|personnes?)', 'Bénéficiaires', 'Beneficiaries'),
]


def _safe_text(value: Any, fallback: str = '') -> str:
    if value is None:
        return fallback
    text = str(value).strip()
    return text if text else fallback


def _clean_spaces(value: Any) -> str:
    return ' '.join(_safe_text(value).replace('\n', ' ').split())


def _sentence(text: str) -> str:
    text = _clean_spaces(text)
    if text and text[-1] not in '.!?':
        text += '.'
    return text


def _parse_float(value: Any) -> float | None:
    try:
        if value is None:
            return None
        if isinstance(value, str):
            value = value.replace(',', '.').strip()
        return float(value)
    except Exception:
        return None


def _format_date(value: Any, lang: str) -> str:
    text = _safe_text(value)
    if not text:
        return datetime.now().strftime('%d/%m/%Y' if lang == 'fr' else '%Y-%m-%d')
    try:
        dt = datetime.fromisoformat(text.replace('Z', '+00:00'))
        return dt.strftime('%d/%m/%Y' if lang == 'fr' else '%Y-%m-%d')
    except Exception:
        return text[:10] if len(text) >= 10 else text


def _rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color.strip().lstrip('#').upper())


def _pil_rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.strip().lstrip('#')
    return tuple(int(h[i:i + 2], 16) for i in (0, 2, 4))


def _norm_lower(text: str) -> str:
    return _clean_spaces(text).lower()


def resource_path(relative: str) -> Path:
    candidates = [
        Path(getattr(sys, '_MEIPASS', Path(__file__).resolve().parent.parent)) / relative,
        Path(__file__).resolve().parent / relative,
        Path(__file__).resolve().parent / relative.replace('desktop_sync_app/', ''),
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return candidates[0]


class PremiumActivityReportBuilder:
    """Editable premium DOCX activity report.

    Page 1 is built with real Word tables, paragraphs and images so the NGO can edit text.
    Maps, photos and pictograms remain images, but all title/narrative/KPI/caption text is editable.
    """

    def __init__(self, base_folder: Path, default_org_name: str = '') -> None:
        self.base_folder = Path(base_folder)
        self.default_org_name = _safe_text(default_org_name)
        self.assets_dir = resource_path('desktop_sync_app/assets')
        self.maps_dir = self.assets_dir / 'maps' / 'naturalearth_lowres'
        self._shape_cache: list[Any] | None = None
        self._icon_cache: dict[str, Path] = {}
        self._temp_dir: Path | None = None

    def build(self, output: Path, project_code: str, project_name: str, records: list[Any], lang: str) -> None:
        lang = 'en' if str(lang).lower().startswith('en') else 'fr'
        if not records:
            return
        data = self._prepare_data(project_code, project_name, records, lang)
        output.parent.mkdir(parents=True, exist_ok=True)
        with tempfile.TemporaryDirectory(prefix='grantproof_report_assets_') as tmp:
            self._temp_dir = Path(tmp)
            self._icon_cache = {}
            hero_size = (1005, 510)
            data['hero_picture'] = self._prepare_photo(data.get('hero_image'), self._temp_dir / 'hero.png', hero_size)
            data['map_picture'] = self._render_map_image(data, self._temp_dir / 'map.png', hero_size, lang)
            data['header_country_picture'] = self._render_country_silhouette(data, self._temp_dir / 'country.png', (260, 170), for_header=True)
            template_path = self._report_template_path()
            if template_path.exists():
                self._build_from_word_template(output, data, lang, template_path)
            else:
                doc = Document()
                self._configure_docx(doc)
                self._build_page_one(doc, data, lang)
                self._build_annex_pages(doc, data, lang)
                doc.save(output)
            self._temp_dir = None
            self._icon_cache = {}


    # ------------------------------------------------------------------ Word template report

    def _report_template_path(self) -> Path:
        return self.assets_dir / 'templates' / 'grantproof_report_template_v1_placeholders.docx'

    def _build_from_word_template(self, output: Path, data: dict[str, Any], lang: str, template_path: Path) -> None:
        """Generate the premium report from the user-provided editable Word template.

        The template is kept as-is: colors, shapes, tables, positioning and fonts are
        preserved. Only placeholders are replaced and a small set of known placeholder
        images are swapped inside word/media. The main photo is prepared with rounded
        corners before being inserted, per the user's instruction.
        """
        doc = Document(str(template_path))
        kpis = self._template_kpis(data, lang)
        sectors = self._template_sectors(data)
        gps = data.get('gps_point')
        gps_text = ''
        if gps:
            lat, lon, _ = gps
            gps_text = f'GPS : {lat:.4f}, {lon:.4f}'
        replacements = {
            '{{REPORT_TITLE}}': data.get('title', ''),
            '{{PROJECT_NAME}}': data.get('project_name', ''),
            '{{PROJECT_CODE}}': data.get('project_code', ''),
            '{{ORGANIZATION}}': data.get('org_name', ''),
            '{{LOCATION}}': data.get('location', ''),
            '{{REPORT_DATE}}': data.get('report_date', ''),
            '{{SUMMARY}}': self._truncate_text(data.get('summary', ''), 300),
            '{{TARGET_GROUP}}': data.get('target_group', ''),
            '{{EVIDENCE_TYPE}}': data.get('evidence_type', ''),
            '{{PROJECT_OUTPUT}}': data.get('activity_output') or ('Output non renseigné' if lang == 'fr' else 'Output not specified'),
            '{{MAP_IMAGE}}': '',
            '{{MAP_LOCATION_LABEL}}': data.get('map_location_label') or data.get('location', ''),
            '{{GPS_COORDINATES}}': gps_text,
            '{{SECTOR_1_LABEL}}': sectors[0].get('label', '') if len(sectors) > 0 else '',
            '{{SECTOR_2_LABEL}}': sectors[1].get('label', '') if len(sectors) > 1 else '',
        }
        for i in range(3):
            item = kpis[i] if i < len(kpis) else {'value': '', 'label': ''}
            replacements[f'{{{{KPI_{i+1}_VALUE}}}}'] = str(item.get('value', ''))
            replacements[f'{{{{KPI_{i+1}_LABEL}}}}'] = str(item.get('label', ''))
        self._replace_template_text(doc, replacements)
        self._fill_template_map(doc, data, lang)
        self._fill_template_annex(doc, data, lang)
        doc.save(output)

        media = {}
        media['image1.png'] = data.get('header_country_picture')
        media['image5.jpg'] = self._prepare_template_photo(
            data.get('hero_image'),
            self._temp_dir / 'template_hero_rounded.jpg',
            self._template_media_size(template_path, 'image5.jpg', (346, 333)),
        )
        media['image18.png'] = self._prepare_template_png(
            data.get('map_picture'),
            self._temp_dir / 'template_map.png',
            self._template_media_size(template_path, 'image18.png', (433, 273)),
            rounded=True,
        )
        for idx, item in enumerate(kpis[:3], start=7):
            media[f'image{idx}.png'] = self._plain_icon(self._kpi_icon_key(item))
        if len(sectors) > 0:
            media['image15.png'] = self._sector_icon('food_security' if sectors[0].get('key') in {'food_security', 'food'} else sectors[0].get('icon', sectors[0].get('key', 'coordination')))
        if len(sectors) > 1:
            media['image16.png'] = self._sector_icon('agriculture' if sectors[1].get('key') == 'agriculture' else sectors[1].get('icon', sectors[1].get('key', 'coordination')))
        self._replace_template_media(output, {name: path for name, path in media.items() if path})

    def _template_kpis(self, data: dict[str, Any], lang: str) -> list[dict[str, Any]]:
        kpis = list(data.get('kpis') or [])
        while len(kpis) < 3:
            kinds = {item.get('kind') for item in kpis}
            if 'media' not in kinds:
                kpis.append({'kind': 'media', 'value': str(data.get('media_count', 0)), 'label': 'Médias liés' if lang == 'fr' else 'Linked media', 'icon': 'media'})
            elif 'evidence' not in kinds:
                kpis.append({'kind': 'evidence', 'value': str(data.get('evidence_count', 1)), 'label': 'Preuve consolidée' if lang == 'fr' else 'Consolidated evidence', 'icon': 'check'})
            else:
                break
        preferred_order = ['beneficiary', 'beneficiaries', 'people', 'farmer', 'kit', 'kits', 'group', 'groups', 'media', 'evidence']
        def rank(item: dict[str, Any]):
            kind = _safe_text(item.get('kind')).lower()
            try:
                priority = preferred_order.index(kind)
            except ValueError:
                priority = 99
            return priority, _safe_text(item.get('label'))
        normalized = []
        seen = set()
        for item in sorted(kpis, key=rank):
            sig = (item.get('kind'), item.get('value'))
            if sig in seen:
                continue
            normalized.append(item)
            seen.add(sig)
            if len(normalized) == 3:
                break
        return normalized

    def _template_sectors(self, data: dict[str, Any]) -> list[dict[str, str]]:
        sectors = list(data.get('sectors') or [])
        if not sectors:
            return [{'key': 'coordination', 'label': 'Coordination', 'icon': 'coordination'}]
        ordered = sorted(sectors, key=lambda s: 0 if s.get('key') in {'food_security', 'food'} else (1 if s.get('key') == 'agriculture' else 2))
        return ordered[:2]

    def _replace_template_text(self, doc: Document, replacements: dict[str, str]) -> None:
        def replace_in_paragraph(paragraph) -> None:
            for run in paragraph.runs:
                text = run.text
                if not text:
                    continue
                for key, value in replacements.items():
                    if key in text:
                        text = text.replace(key, _safe_text(value))
                run.text = text
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)

    def _fill_template_map(self, doc: Document, data: dict[str, Any], lang: str) -> None:
        if len(doc.tables) < 5:
            return
        table = doc.tables[4]
        if len(table.rows) < 2 or len(table.rows[1].cells) < 2:
            return
        map_cell = table.cell(1, 1)
        self._clear_cell_text_only(map_cell)
        try:
            table.rows[1].height = Inches(0.95)
            table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for row_idx in range(2, len(table.rows)):
                table.rows[row_idx].height = Inches(0.05)
                table.rows[row_idx].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                if len(table.rows[row_idx].cells) > 1:
                    self._clear_cell_text_only(table.cell(row_idx, 1))
        except Exception:
            pass
        p = map_cell.paragraphs[0] if map_cell.paragraphs else map_cell.add_paragraph()
        self._pconf(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=2, line=1.0)
        prepared = self._prepare_template_png(data.get('map_picture'), self._temp_dir / 'template_map_inline.png', (420, 240), rounded=True)
        if prepared.exists():
            p.add_run().add_picture(str(prepared), width=Inches(1.80))
        caption = map_cell.add_paragraph()
        self._pconf(caption, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.0)
        self._run(caption, data.get('map_location_label') or data.get('location', ''), bold=True, size=9.0, color=BLUE_2)
        if data.get('gps_point'):
            lat, lon, _ = data['gps_point']
            gps = map_cell.add_paragraph()
            self._pconf(gps, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.0)
            self._run(gps, f'GPS : {lat:.4f}, {lon:.4f}', size=8.5, color=MUTED)

    def _fill_template_annex(self, doc: Document, data: dict[str, Any], lang: str) -> None:
        images = list(data.get('annex_images') or [])[:4]
        if not images:
            self._remove_template_annex(doc)
            return
        if len(doc.tables) < 6:
            return
        table = doc.tables[5]
        for idx in range(4):
            row_idx = 0 if idx < 2 else 1
            col_idx = idx % 2
            if row_idx >= len(table.rows) or col_idx >= len(table.rows[row_idx].cells):
                continue
            cell = table.cell(row_idx, col_idx)
            self._clear_cell_text_only(cell)
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            self._pconf(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, after=2, line=1.0)
            if idx < len(images) and Path(images[idx]).exists():
                prepared = self._prepare_template_png(Path(images[idx]), self._temp_dir / f'annex_{idx+1}.png', (620, 390), rounded=True)
                paragraph.add_run().add_picture(str(prepared), width=Inches(2.85))
                caption = cell.add_paragraph()
                self._pconf(caption, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.0)
                self._run(caption, f"Média complémentaire {idx + 1}" if lang == 'fr' else f"Additional media {idx + 1}", size=9.0, color=MUTED)

    def _remove_template_annex(self, doc: Document) -> None:
        # Keep the user-provided template structure untouched. The annex remains
        # available as a stable editable page even when there are no extra media.
        return

    def _clear_cell_text_only(self, cell) -> None:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.text = ''

    def _template_media_size(self, template_path: Path, filename: str, fallback: tuple[int, int]) -> tuple[int, int]:
        try:
            with zipfile.ZipFile(template_path, 'r') as zf:
                data = zf.read(f'word/media/{filename}')
            tmp = self._temp_dir / f'_probe_{filename}'
            tmp.write_bytes(data)
            with Image.open(tmp) as im:
                return im.size
        except Exception:
            return fallback

    def _prepare_template_photo(self, source: Path | None, output: Path, size: tuple[int, int]) -> Path:
        try:
            if source and Path(source).exists():
                im = Image.open(source).convert('RGB')
                im = self._cover_crop(im, size)
            else:
                im = Image.new('RGB', size, '#F1F5FA')
                d = ImageDraw.Draw(im)
                d.rounded_rectangle((0, 0, size[0] - 1, size[1] - 1), radius=max(18, int(min(size) * 0.08)), fill='#F1F5FA', outline='#D9E4F2', width=3)
                text = 'Photo principale indisponible'
                font = self._font(18, bold=True)
                tw = self._text_w(d, text, font)
                d.text(((size[0] - tw) / 2, size[1] / 2 - 10), text, fill='#6B7280', font=font)
            radius = max(18, int(min(size) * 0.08))
            rgba = im.convert('RGBA')
            mask = Image.new('L', size, 0)
            d = ImageDraw.Draw(mask)
            d.rounded_rectangle((0, 0, size[0] - 1, size[1] - 1), radius=radius, fill=255)
            rounded = Image.new('RGBA', size, (255, 255, 255, 0))
            rounded.alpha_composite(rgba)
            rounded.putalpha(mask)
            border = Image.new('RGBA', size, (255, 255, 255, 0))
            db = ImageDraw.Draw(border)
            db.rounded_rectangle((1, 1, size[0] - 2, size[1] - 2), radius=radius, outline=(220, 228, 238, 255), width=max(2, int(min(size) * 0.01)))
            rounded = Image.alpha_composite(rounded, border)
            flat = Image.new('RGB', size, '#FFFFFF')
            flat.paste(rounded, mask=rounded.getchannel('A'))
            flat.save(output, quality=94)
            return output
        except Exception:
            return output

    def _prepare_template_png(self, source: Path | None, output: Path, size: tuple[int, int], rounded: bool = False) -> Path:
        if source and Path(source).exists():
            try:
                im = Image.open(source).convert('RGBA')
                im = ImageOps.contain(im, size, Image.Resampling.LANCZOS)
                canvas = Image.new('RGBA', size, (255, 255, 255, 0))
                canvas.alpha_composite(im, ((size[0] - im.width) // 2, (size[1] - im.height) // 2))
            except Exception:
                canvas = Image.new('RGBA', size, (255, 255, 255, 0))
        else:
            canvas = Image.new('RGBA', size, (255, 255, 255, 0))
        if rounded:
            radius = max(16, int(min(size) * 0.08))
            mask = Image.new('L', size, 0)
            d = ImageDraw.Draw(mask)
            d.rounded_rectangle((0, 0, size[0] - 1, size[1] - 1), radius=radius, fill=255)
            canvas.putalpha(mask)
        canvas.save(output)
        return output

    def _replace_template_media(self, docx_path: Path, replacements: dict[str, Path]) -> None:
        tmp = docx_path.with_suffix(docx_path.suffix + '.tmp')
        try:
            with zipfile.ZipFile(docx_path, 'r') as zin, zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    content = zin.read(item.filename)
                    if item.filename.startswith('word/media/'):
                        name = Path(item.filename).name
                        replacement = replacements.get(name)
                        if replacement and Path(replacement).exists():
                            content = Path(replacement).read_bytes()
                    zout.writestr(item, content)
            tmp.replace(docx_path)
        except Exception:
            if tmp.exists():
                tmp.unlink()

    # ------------------------------------------------------------------ data

    def _prepare_data(self, project_code: str, project_name: str, records: list[Any], lang: str) -> dict[str, Any]:
        evidence_records = [record for record in records if getattr(record, 'kind', '') == 'evidence']
        story_records = [record for record in records if getattr(record, 'kind', '') == 'story']
        report_mode = 'story' if story_records and not evidence_records else 'evidence'
        primary_record = self._select_primary_record(story_records if report_mode == 'story' else (evidence_records or records))
        project = self._project_payload(records)
        org_name = self._extract_org_name(project, records)
        country_raw = _safe_text(project.get('country') or self._best_raw_value(records, ['country', 'projectCountry']))
        country = self._canonical_country(country_raw)
        country_display = self._display_country(country, country_raw, lang)
        location = self._extract_location(primary_record, records, country_display)
        report_date = _format_date(getattr(primary_record, 'raw', {}).get('createdAt') or getattr(primary_record, 'created_at', ''), lang)
        activity_raw = self._extract_activity(primary_record, records, lang)
        activity = self._polish_activity(activity_raw, lang)
        sectors = self._detect_sectors(records, activity, project_name, lang)
        media_files = self._all_media(records)
        image_files = [path for path in media_files if Path(path).suffix.lower() in IMAGE_EXTENSIONS]
        hero_image = self._select_hero_image(image_files, primary_record, records)
        annex_images = [Path(path) for path in image_files if hero_image is None or Path(path) != Path(hero_image)]
        gps = self._extract_gps(primary_record, records)
        raw = getattr(primary_record, 'raw', {}) or {}
        description_source = getattr(primary_record, 'description', '') or raw.get('description') or raw.get('summary') or raw.get('quote')
        description = self._polish_source_description(_safe_text(description_source))
        subtype = ('Success story' if lang != 'fr' else 'Success story') if report_mode == 'story' else self._evidence_type_label(primary_record, lang)
        title = self._story_report_title(primary_record, activity, location, country_display, lang) if report_mode == 'story' else self._report_title(activity, location, country_display, lang)
        indicators = self._extract_indicators(records, lang)
        if report_mode == 'story':
            indicators = self._story_indicators(primary_record, story_records or records, indicators, lang)
        quantity = indicators[0] if indicators else ({'value': str(len(story_records) or len(records)), 'label': 'Success stories' if lang == 'fr' else 'Success stories', 'icon': 'check', 'kind': 'story'} if report_mode == 'story' else {'value': str(len(evidence_records) or len(records)), 'label': 'Preuves' if lang == 'fr' else 'Evidence items', 'icon': 'check', 'kind': 'evidence'})
        activity_output = self._extract_activity_output(primary_record, records, lang)
        summary = self._story_narrative(primary_record, project_name, activity, description, location, lang) if report_mode == 'story' else self._narrative(activity, project_name, description, indicators, location, sectors, lang)
        highlights = self._story_highlights(primary_record, project_code, activity, location, len(media_files), lang) if report_mode == 'story' else self._highlights(activity, project_code, description, indicators, len(media_files), location, sectors, lang)
        target_group = self._story_target_group(primary_record, description, lang) if report_mode == 'story' else self._target_group_label(quantity, sectors, description, lang)
        return {
            'project_code': project_code,
            'project_name': project_name,
            'org_name': org_name,
            'country': country,
            'country_display': country_display,
            'location': location,
            'report_date': report_date,
            'activity': activity,
            'activity_output': activity_output,
            'title': title,
            'sectors': sectors,
            'quantity': quantity,
            'kpis': indicators[:4],
            'media_count': len(media_files),
            'evidence_count': max(1, len(evidence_records) or len(story_records) or len(records)),
            'hero_image': hero_image,
            'annex_images': annex_images,
            'gps_point': gps,
            'description': description,
            'evidence_type': subtype,
            'summary': summary,
            'highlights': highlights,
            'target_group': target_group,
            'map_location_label': gps[2] if gps else country_display,
            'report_mode': report_mode,
        }

    def _patch_page_margins_for_header_bleed(self, output: Path) -> None:
        """Apply a narrow OOXML margin patch after python-docx save.

        python-docx refuses negative page margins, while LibreOffice/Word PDF export can
        keep an unavoidable printable-area offset even when margins are zero. A small
        negative left/right margin lets the first header table visually bleed to the
        page edges, while all body tables remain manually centered at the requested
        2cm safe width.
        """
        try:
            tmp = output.with_suffix(output.suffix + '.tmp')
            with zipfile.ZipFile(output, 'r') as zin, zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    content = zin.read(item.filename)
                    if item.filename == 'word/document.xml':
                        xml = content.decode('utf-8')
                        xml = re.sub(r'(<w:pgMar\b[^>]*\bw:left=")[^"]*(")', r'\1-1138\2', xml)
                        xml = re.sub(r'(<w:pgMar\b[^>]*\bw:right=")[^"]*(")', r'\1-1138\2', xml)
                        content = xml.encode('utf-8')
                    zout.writestr(item, content)
            tmp.replace(output)
        except Exception:
            # The report must remain generatable even if an external zip patch fails.
            pass

    def _configure_docx(self, doc: Document) -> None:
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(14)
        section.page_height = Inches(8.5)
        # Header must be full-bleed at the top/left/right. Body spacing is created with centered tables.
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0.08)
        # Body pages use 2cm left/right margins. The header blue background is a separate
        # full-bleed shape so it is not constrained by these body margins.
        section.left_margin = Inches(0.79)
        section.right_margin = Inches(0.79)
        section.header_distance = Inches(0)
        section.footer_distance = Inches(0)
        styles = doc.styles
        styles['Normal'].font.name = 'Arial'
        styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        styles['Normal'].font.size = Pt(10)
        styles['Normal'].font.color.rgb = _rgb(DARK)

    # ------------------------------------------------------------------ editable premium dashboard page

    def _build_page_one(self, doc: Document, data: dict[str, Any], lang: str) -> None:
        self._build_header(doc, data, lang)

        # Real spacer columns create visible breathing room between editable Word cards.
        top = doc.add_table(rows=1, cols=5)
        self._table_no_borders(top)
        top.alignment = WD_TABLE_ALIGNMENT.CENTER
        top.autofit = False
        self._set_row_height(top.rows[0], Inches(4.18))
        # Body width is kept to 14in - 2cm - 2cm = approx. 12.43in.
        # The header remains true full-bleed because the section margins stay at zero.
        top_widths = [4.05, 0.24, 4.05, 0.24, 3.85]
        self._set_table_grid(top, top_widths)
        card_specs = {
            0: (4.05, 22, 18, 36, 36),
            2: (4.05, 22, 18, 26, 26),
            4: (3.85, 18, 16, 22, 22),
        }
        for idx, cell in enumerate(top.rows[0].cells):
            self._set_cell_width(cell, Inches(top_widths[idx]))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if idx in card_specs:
                width, top_pad, bottom_pad, start_pad, end_pad = card_specs[idx]
                self._set_cell_width(cell, Inches(width))
                self._cell_margins(cell, top=top_pad, bottom=bottom_pad, start=start_pad, end=end_pad)
                self._shade_cell(cell, WHITE)
                if idx != 4:
                    self._border_cell(cell, BORDER_LIGHT)
            else:
                self._cell_margins(cell, top=0, bottom=0, start=0, end=0)
        self._build_overview(top.cell(0, 0), data, lang)
        self._build_kpis(top.cell(0, 2), data, lang)
        self._build_media_sector(top.cell(0, 4), data, lang)

        self._spacer(doc, 0)
        bottom = doc.add_table(rows=1, cols=3)
        self._table_no_borders(bottom)
        bottom.alignment = WD_TABLE_ALIGNMENT.CENTER
        bottom.autofit = False
        self._set_row_height(bottom.rows[0], Inches(2.76))
        bottom_widths = [5.84, 0.25, 6.34]
        self._set_table_grid(bottom, bottom_widths)
        for idx, cell in enumerate(bottom.rows[0].cells):
            self._set_cell_width(cell, Inches(bottom_widths[idx]))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if idx in {0, 2}:
                self._cell_margins(cell, top=18, bottom=16, start=42, end=42)
                self._shade_cell(cell, SOFT_BLUE)
                self._border_cell(cell, BORDER_LIGHT)
            else:
                self._cell_margins(cell, top=0, bottom=0, start=0, end=0)
        self._build_highlights(bottom.cell(0, 0), data, lang)
        self._build_location(bottom.cell(0, 2), data, lang)

    def _add_header_bleed_background(self, paragraph) -> None:
        """Add an editable-DOCX-safe full-bleed blue rectangle behind the header text.

        This is not a full-page image: it is a small vector/VML rectangle used only
        as the header background. All report text and tables remain editable.
        """
        try:
            shape = parse_xml(
                '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'xmlns:v="urn:schemas-microsoft-com:vml" '
                'xmlns:o="urn:schemas-microsoft-com:office:office">'
                '<w:pict>'
                '<v:rect id="GrantProofHeaderBleed" stroked="f" fillcolor="#%s" '
                'style="position:absolute;left:-0.79in;top:0;width:14in;height:1.18in;'
                'z-index:-251654144;mso-position-horizontal-relative:page;'
                'mso-position-vertical-relative:page;mso-wrap-style:none">'
                '<v:fill color="#%s"/>'
                '</v:rect>'
                '</w:pict>'
                '</w:r>' % (BLUE, BLUE)
            )
            paragraph._p.insert(0, shape)
        except Exception:
            pass

    def _build_header(self, doc: Document, data: dict[str, Any], lang: str) -> None:
        table = doc.add_table(rows=1, cols=4)
        self._table_no_borders(table)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        self._set_row_height(table.rows[0], Inches(1.18))
        # Header text grid stays inside the 2cm safe area. A separate full-bleed
        # blue shape behind it reaches the page edges.
        widths = [8.25, 1.00, 0.05, 3.13]
        self._set_table_grid(table, widths)
        for idx, (cell, width) in enumerate(zip(table.rows[0].cells, widths)):
            self._set_cell_width(cell, Inches(width))
            self._shade_cell(cell, BLUE)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if idx == 0:
                self._cell_margins(cell, top=20, bottom=16, start=90, end=42)
            elif idx == 1:
                self._cell_margins(cell, top=10, bottom=10, start=0, end=0)
            elif idx == 2:
                self._cell_margins(cell, top=14, bottom=14, start=0, end=0)
            else:
                # Meta block content is vertically centered and shifted about 1.5cm to the right.
                self._cell_margins(cell, top=18, bottom=16, start=850, end=35)
        title_cell, map_cell, sep_cell, meta_cell = table.rows[0].cells
        self._shade_cell(sep_cell, 'D7E6F7')
        p = self._cell_p(title_cell)
        self._add_header_bleed_background(p)
        self._pconf(p, after=4, line=0.98)
        title = data['title'].upper()
        title_size = 18.0 if len(title) < 58 else (16.0 if len(title) < 82 else 14.2)
        self._run(p, title, bold=True, size=title_size, color=WHITE, font='Arial')
        p = title_cell.add_paragraph()
        self._pconf(p, after=0, line=1.02)
        prefix = 'Projet' if lang == 'fr' else 'Project'
        self._run(p, f"{prefix} : {data['project_name']} ({data['project_code']})", size=11.0, color=WHITE, font='Arial')
        p = self._cell_p(map_cell)
        self._pconf(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
        if Path(data['header_country_picture']).exists():
            p.add_run().add_picture(str(data['header_country_picture']), width=Inches(0.98))
        meta = meta_cell.add_table(rows=3, cols=2)
        self._table_no_borders(meta)
        meta.autofit = False
        self._set_table_grid(meta, [0.30, 2.10])
        values = [
            ('meta_org', self._truncate_text(data['org_name'], 32)),
            ('meta_pin', self._truncate_text(data['location'], 36)),
            ('meta_calendar', data['report_date']),
        ]
        for i, (icon, value) in enumerate(values):
            row = meta.rows[i]
            self._set_row_height(row, Inches(0.265))
            ico = row.cells[0]; txt = row.cells[1]
            self._set_cell_width(ico, Inches(0.30)); self._set_cell_width(txt, Inches(2.10))
            self._cell_margins(ico, top=0, bottom=0, start=0, end=10)
            self._cell_margins(txt, top=0, bottom=0, start=0, end=0)
            ico.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            txt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = self._cell_p(ico)
            self._pconf(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.0)
            p.add_run().add_picture(str(self._white_icon(icon)), width=Inches(0.16))
            p = self._cell_p(txt)
            self._pconf(p, after=0, line=1.0)
            self._run(p, value, size=10.0, color=WHITE, font='Arial')

    def _build_overview(self, cell, data: dict[str, Any], lang: str) -> None:
        self._shade_cell(cell, WHITE)
        self._border_cell(cell, BORDER_LIGHT)
        self._section_title(cell, 'APERÇU DU PROJET' if lang == 'fr' else 'PROJECT OVERVIEW', 'model_section_overview', underline_width=2.65)
        p = cell.add_paragraph()
        summary_size = 10.0
        self._pconf(p, before=4, line=1.03, after=4)
        self._run(p, data['summary'], size=summary_size, color=DARK)
        # User request: move the Lieu / Public cible / Type de preuve block down by about 2cm.
        self._spacer_cell(cell, 67)
        info = cell.add_table(rows=3, cols=2)
        info.alignment = WD_TABLE_ALIGNMENT.LEFT
        self._table_no_borders(info)
        info.autofit = False
        self._set_table_grid(info, [0.34, 3.24])
        rows = [
            ('model_info_location', 'Lieu' if lang == 'fr' else 'Location', data['location']),
            ('model_info_people', 'Public cible' if lang == 'fr' else 'Target group', data['target_group']),
            ('model_info_photo', 'Type de preuve' if lang == 'fr' else 'Evidence type', data['evidence_type']),
        ]
        for row, (icon, label, value) in zip(info.rows, rows):
            ico, txt = row.cells
            self._set_cell_width(ico, Inches(0.34))
            self._set_cell_width(txt, Inches(3.24))
            self._cell_margins(ico, top=6, bottom=6, start=0, end=14)
            self._cell_margins(txt, top=6, bottom=6, start=0, end=0)
            ico.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            txt.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = self._cell_p(ico)
            self._pconf(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
            p.add_run().add_picture(str(self._plain_icon(icon)), width=Inches(0.24))
            p = self._cell_p(txt)
            self._pconf(p, line=1.03, after=0)
            self._run(p, label, bold=True, size=10.0, color=BLUE_2)
            p = txt.add_paragraph()
            self._pconf(p, line=1.03, after=0)
            self._run(p, value, size=10.0, color=DARK)

    def _kpi_icon_key(self, kpi: dict[str, Any]) -> str:
        kind = _safe_text(kpi.get('kind')).lower()
        icon = _safe_text(kpi.get('icon')).lower()
        joined = f"{kind} {icon} {_safe_text(kpi.get('label')).lower()}"
        if any(term in joined for term in ['beneficiary', 'bénéficiaire', 'people', 'person']):
            return 'model_kpi_beneficiaries'
        if any(term in joined for term in ['semence', 'intrant', 'seed']):
            return 'distribution_semences'
        if any(term in joined for term in ['kit', 'kits', 'box']):
            return 'distribution_semences'
        if any(term in joined for term in ['group', 'groupement']):
            return 'model_kpi_groups'
        if any(term in joined for term in ['media', 'média', 'photo', 'camera']):
            return 'model_kpi_media'
        if any(term in joined for term in ['evidence', 'preuve', 'check']):
            return 'model_kpi_evidence'
        if any(term in joined for term in ['cash', 'voucher']):
            return 'cash_assistance'
        if any(term in joined for term in ['shelter', 'abri']):
            return 'shelter'
        if any(term in joined for term in ['training', 'formation']):
            return 'formation_atelier'
        return 'model_kpi_evidence'

    def _build_kpis(self, cell, data: dict[str, Any], lang: str) -> None:
        self._section_title(cell, 'CHIFFRES CLÉS' if lang == 'fr' else 'KEY FIGURES', 'model_section_kpis', underline_width=4.2)
        kpis = list(data.get('kpis') or [])
        normalized = []
        seen = set()
        preferred_order = ['beneficiary', 'beneficiaries', 'people', 'media', 'evidence', 'kit', 'kits', 'group', 'groups']
        while len(kpis) < 3:
            kinds = {item.get('kind') for item in kpis}
            if 'media' not in kinds:
                kpis.append({'kind': 'media', 'value': str(data['media_count']), 'label': 'Médias liés' if lang == 'fr' else 'Linked media', 'icon': 'media'})
            elif 'evidence' not in kinds:
                kpis.append({'kind': 'evidence', 'value': str(data['evidence_count']), 'label': 'Preuve consolidée' if lang == 'fr' else 'Consolidated evidence', 'icon': 'check'})
            else:
                break
        def rank(item):
            kind = _safe_text(item.get('kind')).lower()
            try:
                priority = preferred_order.index(kind)
            except ValueError:
                priority = 99
            return priority, -len(str(item.get('value', '')))
        for item in sorted(kpis, key=rank):
            sig = (item.get('kind'), item.get('value'))
            if sig in seen:
                continue
            normalized.append(item)
            seen.add(sig)
            if len(normalized) == 3:
                break
        grid = cell.add_table(rows=1, cols=3)
        self._table_no_borders(grid)
        grid.autofit = False
        self._set_table_grid(grid, [1.24, 1.24, 1.24])
        self._set_row_height(grid.rows[0], Inches(2.30))
        for kcell in grid.rows[0].cells:
            self._set_cell_width(kcell, Inches(1.24))
            self._cell_margins(kcell, top=14, bottom=12, start=14, end=14)
            self._shade_cell(kcell, WHITE)
            # No blue frame around the KPI mini-cards.
            kcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for idx, kpi in enumerate(normalized[:3]):
            kcell = grid.cell(0, idx)
            p = self._cell_p(kcell)
            self._pconf(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
            p.add_run().add_picture(str(self._plain_icon(self._kpi_icon_key(kpi))), width=Inches(0.64))
            p = kcell.add_paragraph()
            self._pconf(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=0.95)
            self._run(p, self._truncate_text(str(kpi.get('value', '')), 14), bold=True, size=26.0, color=BLUE_2, font='Arial')
            p = kcell.add_paragraph()
            self._pconf(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=0.8)
            self._run(p, '—', bold=True, size=14.0, color=ORANGE, font='Arial')
            p = kcell.add_paragraph()
            self._pconf(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line=1.0)
            self._run(p, self._truncate_text(str(kpi.get('label', '')), 24), size=10.0, color=DARK, font='Arial')

    def _build_project_logic(self, cell, data: dict[str, Any], lang: str) -> None:
        activity = _safe_text(data.get('activity'))
        output = _safe_text(data.get('activity_output'))
        if not activity and not output:
            return
        self._spacer_cell(cell, 5)
        title = 'INTÉGRATION AU CADRE LOGIQUE DU PROJET' if lang == 'fr' else 'PROJECT LOGFRAME INTEGRATION'
        self._section_title(cell, title, 'model_section_logframe', underline_width=3.2)
        if activity:
            p = cell.add_paragraph()
            self._pconf(p, after=1, line=1.05)
            self._run(p, 'Activité : ', bold=True, size=8.8, color=ORANGE)
            self._run(p, activity, size=8.8)
        if output:
            p = cell.add_paragraph()
            self._pconf(p, after=0, line=1.05)
            self._run(p, 'Output : ', bold=True, size=8.8, color=ORANGE)
            self._run(p, output, size=8.8)

    def _build_media_sector(self, cell, data: dict[str, Any], lang: str) -> None:
        p = self._cell_p(cell)
        self._pconf(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
        if Path(data['hero_picture']).exists():
            p.add_run().add_picture(str(data['hero_picture']), width=Inches(3.38), height=Inches(2.18))
        self._spacer_cell(cell, 4)
        self._section_title(cell, 'ALIGNEMENT SECTORIEL' if lang == 'fr' else 'SECTOR ALIGNMENT', 'model_section_sector', underline_width=2.6)
        self._spacer_cell(cell, 4)
        sectors = sorted(data['sectors'][:2], key=lambda s: 0 if s.get('key') in {'food_security','food'} else (1 if s.get('key') == 'agriculture' else 2)) or [{'key': 'coordination', 'label': 'Coordination', 'icon': 'coordination'}]
        table = cell.add_table(rows=1, cols=max(1, len(sectors)))
        self._table_no_borders(table)
        table.autofit = False
        self._set_table_grid(table, [1.58 for _ in sectors])
        self._set_row_height(table.rows[0], Inches(0.96))
        for scell, sector in zip(table.rows[0].cells, sectors):
            self._set_cell_width(scell, Inches(1.58))
            self._cell_margins(scell, top=4, bottom=6, start=6, end=6)
            scell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = self._cell_p(scell)
            self._pconf(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=2, line=1.0)
            icon_key = 'food_security' if sector.get('key') in {'food_security', 'food'} else ('agriculture' if sector.get('key') == 'agriculture' else sector['icon'])
            p.add_run().add_picture(str(self._sector_icon(icon_key)), width=Inches(0.62))
            p = scell.add_paragraph()
            self._pconf(p, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0, after=0)
            self._run(p, sector['label'], size=10.0)
        # Real DOCX spacer: keeps the top dashboard band visually full-height in Word/LibreOffice.
        self._spacer_cell(cell, 72)

    def _build_highlights(self, cell, data: dict[str, Any], lang: str) -> None:
        self._section_title(cell, 'FAITS SAILLANTS' if lang == 'fr' else 'HIGHLIGHTS', 'model_section_highlights', underline_width=3.5)
        for item in data['highlights'][:3]:
            p = cell.add_paragraph()
            self._pconf(p, line=1.10, after=3)
            p.paragraph_format.left_indent = Inches(0.20)
            p.paragraph_format.first_line_indent = Inches(-0.16)
            self._run(p, '• ', bold=True, color=BLUE_2, size=10.0)
            self._run(p, item, size=10.0)

    def _build_location(self, cell, data: dict[str, Any], lang: str) -> None:
        self._section_title(cell, 'LOCALISATION' if lang == 'fr' else 'LOCATION', 'model_section_location', underline_width=2.05)
        layout = cell.add_table(rows=1, cols=2)
        self._table_no_borders(layout)
        layout.autofit = False
        self._set_table_grid(layout, [1.95, 4.12])
        self._set_row_height(layout.rows[0], Inches(2.08))
        left, right = layout.rows[0].cells
        self._set_cell_width(left, Inches(1.95))
        self._set_cell_width(right, Inches(4.12))
        self._cell_margins(left, top=8, bottom=0, start=0, end=20)
        self._cell_margins(right, top=0, bottom=0, start=0, end=0)
        left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = self._cell_p(left)
        self._pconf(p, before=0, after=2)
        self._run(p, 'Localisation de l’activité' if lang == 'fr' else 'Activity location', bold=True, size=10.0, color=BLUE_2)
        p = left.add_paragraph()
        self._pconf(p, after=1, line=0.55)
        self._run(p, '────', bold=True, size=6.0, color=ORANGE)
        p = left.add_paragraph()
        self._pconf(p, after=0, line=1.05)
        self._run(p, data['location'], size=10.0, color=DARK)
        p = left.add_paragraph()
        self._pconf(p, after=0, line=1.05)
        if data.get('gps_point'):
            lat, lon, _ = data['gps_point']
            self._run(p, f'GPS : {lat:.4f}, {lon:.4f}', size=10.0, color=MUTED)
        else:
            self._run(p, 'Carte nationale affichée par défaut.', size=10.0, color=MUTED)
        p = self._cell_p(right)
        self._pconf(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
        p.add_run().add_picture(str(data['map_picture']), width=Inches(3.82), height=Inches(1.54))
        p = right.add_paragraph()
        self._pconf(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
        caption = data['map_location_label'] if data.get('gps_point') else data['country_display']
        self._run(p, caption, bold=True, size=10.0, color=BLUE_2)
        # Real DOCX spacer: avoids a half-page report when the evidence summary is short.
        self._spacer_cell(cell, 76)

    def _build_annex_pages(self, doc: Document, data: dict[str, Any], lang: str) -> None:
        images = data.get('annex_images') or []
        if not images:
            return
        for start in range(0, len(images), 6):
            doc.add_page_break()
            self._spacer(doc, 42)
            title = 'ANNEXE VISUELLE — Médias complémentaires' if lang == 'fr' else 'VISUAL ANNEX — Additional media'
            p = doc.add_paragraph()
            self._pconf(p, after=2)
            p.paragraph_format.left_indent = Inches(0.79)
            p.paragraph_format.right_indent = Inches(0.79)
            self._run(p, title, bold=True, size=16, color=BLUE)
            p = doc.add_paragraph()
            self._pconf(p, after=6)
            p.paragraph_format.left_indent = Inches(0.79)
            p.paragraph_format.right_indent = Inches(0.79)
            subtitle = 'Images complémentaires liées à l’activité et conservées dans le dossier technique de preuve.' if lang == 'fr' else 'Additional images linked to the activity and kept in the technical evidence folder.'
            self._run(p, subtitle, size=8.4, color=MUTED)
            chunk = images[start:start + 6]
            cols = min(3, max(1, len(chunk)))
            rows = math.ceil(len(chunk) / cols)
            grid = doc.add_table(rows=rows, cols=cols)
            self._table_no_borders(grid)
            grid.alignment = WD_TABLE_ALIGNMENT.CENTER
            grid.autofit = False
            card_width = 9.60 / cols
            card_height = 2.55 if rows > 1 else 4.25
            for row in grid.rows:
                self._set_row_height(row, Inches(card_height))
                for cell in row.cells:
                    self._set_cell_width(cell, Inches(card_width))
                    self._shade_cell(cell, SOFT_BLUE)
                    self._border_cell(cell, BORDER_LIGHT)
                    self._cell_margins(cell, top=55, bottom=35, start=55, end=55)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for idx, path in enumerate(chunk):
                c = grid.cell(idx // cols, idx % cols)
                pic_px = (880, 560) if cols <= 2 else (720, 450)
                picture = self._prepare_photo(path, self._temp_dir / f'annex_{start + idx}.png', pic_px, fit='cover') if self._temp_dir else None
                p = self._cell_p(c)
                self._pconf(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
                if picture and Path(picture).exists():
                    p.add_run().add_picture(str(picture), width=Inches(min(card_width - 0.45, 4.45)))
                p = c.add_paragraph()
                self._pconf(p, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0, after=0)
                label = f"Média complémentaire {start + idx + 1}" if lang == 'fr' else f'Additional media {start + idx + 1}'
                self._run(p, label, bold=True, size=8.2, color=BLUE_2)
                p = c.add_paragraph()
                self._pconf(p, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.0, after=0)
                self._run(p, f"{data['activity']} | {data['location']} | {data['report_date']}", size=8.0, color=MUTED)

    # ------------------------------------------------------------------ images, icons, maps

    def _icon(self, kind: str) -> Path:
        return self._make_icon(kind, sector=False)

    def _sector_icon(self, kind: str) -> Path:
        return self._make_icon(kind, sector=True)

    def _soft_icon(self, kind: str) -> Path:
        return self._make_icon(kind, sector=False, style='soft')

    def _badge_icon(self, kind: str) -> Path:
        return self._make_icon(kind, sector=False, style='badge')

    def _plain_icon(self, kind: str) -> Path:
        return self._make_icon(kind, sector=False, style='plain')

    def _white_icon(self, kind: str) -> Path:
        return self._make_icon(kind, sector=False, style='white')

    def _resolve_icon_key(self, kind: str, context: str = '') -> str:
        raw = _norm_lower(kind)
        if raw in REPORT_ICON_MAP:
            return REPORT_ICON_MAP[raw]
        if context:
            corpus = _norm_lower(f'{kind} {context}')
            for icon_key, keywords in ACTIVITY_ICON_KEYWORDS:
                if any(keyword in corpus for keyword in keywords):
                    return icon_key
        return REPORT_ICON_MAP.get(raw.replace('-', '_').replace(' ', '_'), 'activity')

    def _report_icon_asset(self, kind: str, style: str = 'plain') -> Path | None:
        if self._temp_dir is None:
            return None
        icon_key = self._resolve_icon_key(kind)
        source = self.assets_dir / 'report_icons' / f'{icon_key}.png'
        # Strict R27 rule: dynamic report pictograms must come from the supplied V3 library only.
        # The only exception is the fixed model_section_* set used for major section titles.
        if not source.exists():
            source = self.assets_dir / 'report_icons' / 'rapport_document.png'
        if not source.exists():
            return None
        cache_key = f'report_{style}_{icon_key}'
        if cache_key in self._icon_cache and self._icon_cache[cache_key].exists():
            return self._icon_cache[cache_key]
        out = self._temp_dir / f'{cache_key}.png'
        im = Image.open(source).convert('RGBA')
        bbox = im.getchannel('A').getbbox()
        if bbox:
            im = im.crop(bbox)
        canvas = Image.new('RGBA', (256, 256), (255, 255, 255, 0))
        max_box = (232, 232) if icon_key.startswith('model_section_') else ((190, 190) if style == 'badge' else (216, 216))
        if im.width and im.height:
            scale = min(max_box[0] / im.width, max_box[1] / im.height)
            new_size = (max(1, int(im.width * scale)), max(1, int(im.height * scale)))
            im = im.resize(new_size, Image.Resampling.LANCZOS)
        # R28: report icons are rendered monochrome.
        # Dynamic V3 pictograms and fixed section-title icons keep their transparent background,
        # but all visible pixels are recolored to a single institutional tone.
        pixels = im.load()
        target = (255, 255, 255) if style in {'white', 'badge'} else _pil_rgb(BLUE_2)
        for y in range(im.height):
            for x in range(im.width):
                r, g, b, a = pixels[x, y]
                if a:
                    pixels[x, y] = (*target, a)
        if style == 'badge':
            badge = Image.new('RGBA', (256, 256), (255, 255, 255, 0))
            d = ImageDraw.Draw(badge)
            d.ellipse((16, 16, 240, 240), fill=(*_pil_rgb(BLUE_2), 255))
            badge.alpha_composite(im, ((256 - im.width) // 2, (256 - im.height) // 2))
            badge.save(out)
        else:
            canvas.alpha_composite(im, ((256 - im.width) // 2, (256 - im.height) // 2))
            canvas.save(out)
        self._icon_cache[cache_key] = out
        return out

    def _humanitarian_icon_asset(self, kind: str, style: str = 'plain') -> Path | None:
        # Backward compatibility: humanitarian icons are now copied/normalized in report_icons.
        return self._report_icon_asset(kind, style)

    def _make_icon(self, kind: str, sector: bool = False, style: str = 'soft') -> Path:
        key = f"{'sector_' if sector else ''}{style}_{kind}"
        if key in self._icon_cache and self._icon_cache[key].exists():
            return self._icon_cache[key]
        asset_style = 'plain'
        if style == 'badge':
            asset_style = 'badge'
        elif style == 'white':
            asset_style = 'white'
        asset_icon = self._report_icon_asset(kind, asset_style)
        if asset_icon:
            self._icon_cache[key] = asset_icon
            return asset_icon
        fallback = self._report_icon_asset('rapport_document', asset_style)
        if fallback:
            self._icon_cache[key] = fallback
            return fallback
        if self._temp_dir is None:
            raise RuntimeError('temporary directory not initialized')
        path = self._temp_dir / f'{key}.png'
        size = 256
        scale = 4
        base = size * scale
        im = Image.new('RGBA', (base, base), (255, 255, 255, 0))
        d = ImageDraw.Draw(im)
        if style == 'badge':
            d.ellipse((22*scale, 22*scale, base - 22*scale, base - 22*scale), fill=(*_pil_rgb(BLUE_2), 255))
            self._draw_line_icon_pil(d, kind, (base // 2, base // 2), (255, 255, 255, 255), scale=4.0)
        elif style == 'white':
            self._draw_line_icon_pil(d, kind, (base // 2, base // 2), (255, 255, 255, 255), scale=3.4)
        else:
            self._draw_line_icon_pil(d, kind, (base // 2, base // 2), (*_pil_rgb(BLUE_2), 255), scale=3.6)
        out = im.resize((size, size), Image.Resampling.LANCZOS)
        out.save(path)
        self._icon_cache[key] = path
        return path

    def _prepare_photo(self, source: Path | None, output: Path, size: tuple[int, int], fit: str = 'cover') -> Path | None:
        def rounded_rgba(rgb: Image.Image, radius: int = 24) -> Image.Image:
            rgba = rgb.convert('RGBA')
            mask = Image.new('L', rgba.size, 0)
            d = ImageDraw.Draw(mask)
            d.rounded_rectangle((0, 0, rgba.size[0] - 1, rgba.size[1] - 1), radius=radius, fill=255)
            rgba.putalpha(mask)
            return rgba
        if source and Path(source).exists():
            try:
                im = Image.open(source).convert('RGB')
                im = self._contain_pad(im, size) if fit == 'contain' else self._cover_crop(im, size)
                rounded_rgba(im, 24).save(output)
                return output
            except Exception:
                pass
        im = Image.new('RGB', size, '#F1F5FA')
        d = ImageDraw.Draw(im)
        d.rounded_rectangle((0, 0, size[0] - 1, size[1] - 1), radius=24, fill='#F1F5FA', outline='#D9E4F2')
        font = self._font(22, bold=True)
        text = 'Photo principale indisponible'
        tw = self._text_w(d, text, font)
        d.text(((size[0] - tw) / 2, size[1] / 2 - 10), text, fill='#6B7280', font=font)
        rounded_rgba(im, 24).save(output)
        return output

    def _contain_pad(self, img: Image.Image, size: tuple[int, int]) -> Image.Image:
        contained = ImageOps.contain(img.convert('RGB'), size, Image.Resampling.LANCZOS)
        canvas = Image.new('RGB', size, '#FFFFFF')
        x = (size[0] - contained.width) // 2
        y = (size[1] - contained.height) // 2
        canvas.paste(contained, (x, y))
        d = ImageDraw.Draw(canvas)
        d.rounded_rectangle((0, 0, size[0] - 1, size[1] - 1), radius=18, outline='#D9E4F2', width=3)
        return canvas

    def _cover_crop(self, img: Image.Image, size: tuple[int, int]) -> Image.Image:
        tw, th = size
        ratio = tw / th
        img = img.convert('RGB')
        if img.width / max(img.height, 1) > ratio:
            new_w = int(img.height * ratio)
            x = (img.width - new_w) // 2
            img = img.crop((x, 0, x + new_w, img.height))
        else:
            new_h = int(img.width / ratio)
            y = max(0, int((img.height - new_h) * 0.35))
            img = img.crop((0, y, img.width, y + new_h))
        return img.resize((tw, th), Image.Resampling.LANCZOS)

    def _render_country_silhouette(self, data: dict[str, Any], output: Path, size: tuple[int, int], for_header: bool = False) -> Path:
        w, h = size
        im = Image.new('RGBA', size, (255, 255, 255, 0))
        d = ImageDraw.Draw(im)
        shape = self._country_shape(data['country'])
        if shape:
            polygons, bbox = self._shape_polygons(shape, data['country'])
            mapper = self._map_transform(bbox, w, h, 10)
            for poly in polygons:
                pts = [mapper(lon, lat) for lon, lat in poly]
                if len(pts) >= 3:
                    d.polygon(pts, fill=(255, 255, 255, 238), outline=(255, 255, 255, 255))
            point = data.get('gps_point')
            if point:
                lat, lon, _ = point
                if self._point_in_bbox(lon, lat, bbox):
                    px, py = mapper(lon, lat)
                    self._draw_pin_pil(d, px, py, 10)
        im.save(output)
        return output

    def _render_map_image(self, data: dict[str, Any], output: Path, size: tuple[int, int], lang: str) -> Path:
        width, height = size
        im = Image.new('RGB', size, '#F8FAFC')
        draw = ImageDraw.Draw(im)
        shape = self._country_shape(data['country'])
        if not shape:
            draw.rounded_rectangle((1, 1, width - 2, height - 2), radius=18, outline='#E2E8F0', fill='#F8FAFC', width=2)
            self._center_text(draw, data['country_display'], (0, 0, width, height), self._font(30, bold=True), '#6B7280')
            im.save(output, quality=95)
            return output
        polygons, bbox = self._shape_polygons(shape, data['country'])
        minx, miny, maxx, maxy = bbox
        dx = max(maxx - minx, 0.01)
        dy = max(maxy - miny, 0.01)
        bbox_expanded = (minx - dx * 0.22, miny - dy * 0.22, maxx + dx * 0.22, maxy + dy * 0.22)
        mapper = self._map_transform(bbox_expanded, width, height, 26)
        target_name = self._shape_country_name(data['country'])
        # background countries
        for sr in self._iter_shapes():
            shp = sr.shape
            sx1, sy1, sx2, sy2 = [float(v) for v in shp.bbox]
            if sx2 < bbox_expanded[0] or sx1 > bbox_expanded[2] or sy2 < bbox_expanded[1] or sy1 > bbox_expanded[3]:
                continue
            name = self._record_name(sr.record)
            is_target = name == target_name
            s_polys, _ = self._shape_polygons(shp, name)
            for poly in s_polys:
                pts = [mapper(lon, lat) for lon, lat in poly]
                if len(pts) >= 3:
                    draw.polygon(pts, fill='#FFFFFF' if is_target else '#F2F4F7', outline='#D9DFE8' if is_target else '#E1E7EF')
        country_font = self._font(30, bold=True)
        cx, cy = mapper((minx + maxx) / 2, (miny + maxy) / 2)
        label = data['country_display'].upper()
        draw.text((cx - self._text_w(draw, label, country_font) / 2, cy - 18), label, fill='#5F6B7A', font=country_font)
        label_font = self._font(14)
        labels = []
        for sr in self._iter_shapes():
            name = self._record_name(sr.record)
            if name == target_name:
                continue
            sx1, sy1, sx2, sy2 = [float(v) for v in sr.shape.bbox]
            if sx2 < bbox_expanded[0] or sx1 > bbox_expanded[2] or sy2 < bbox_expanded[1] or sy1 > bbox_expanded[3]:
                continue
            labels.append((name, (sx1 + sx2) / 2, (sy1 + sy2) / 2))
        for name, lon, lat in labels[:5]:
            lx, ly = mapper(lon, lat)
            txt = self._display_country(name, name, lang).upper()
            draw.text((lx - self._text_w(draw, txt, label_font) / 2, ly - 8), txt, fill='#A1AAB6', font=label_font)
        point = data.get('gps_point')
        if point:
            lat, lon, pin_label = point
            if self._point_in_bbox(lon, lat, bbox_expanded):
                px, py = mapper(lon, lat)
                self._draw_pin_pil(draw, px, py, 25)
                draw.text((px + 18, py - 14), self._truncate_text(pin_label or data['location'], 20), fill='#1D2D44', font=self._font(15, bold=True))
        draw.rounded_rectangle((2, 2, width - 3, height - 3), radius=22, outline='#E1E7EF', width=2)
        im.save(output, quality=95)
        return output

    # ------------------------------------------------------------------ extraction helpers
    # ------------------------------------------------------------------ extraction helpers

    def _project_payload(self, records: list[Any]) -> dict[str, Any]:
        for record in records:
            raw = getattr(record, 'raw', {}) or {}
            project = raw.get('project')
            if isinstance(project, dict) and project:
                return project
        return {}

    def _extract_org_name(self, project: dict[str, Any], records: list[Any]) -> str:
        keys = [
            'organizationName', 'organisationName', 'organization_name', 'organisation_name', 'orgName', 'org_name',
            'ngoName', 'ngo_name', 'ongName', 'ong_name', 'clientName', 'client_name', 'clientOrganizationName',
            'implementingPartner', 'implementing_partner', 'partnerName', 'partner_name', 'agencyName', 'agency_name',
            'workspaceName', 'workspace_name', 'organization', 'organisation', 'ngo', 'ong', 'accountName',
        ]
        for source in [project, *[(getattr(record, 'raw', {}) or {}) for record in records]]:
            for key in keys:
                value = source.get(key) if isinstance(source, dict) else None
                if isinstance(value, dict):
                    value = value.get('name') or value.get('displayName') or value.get('label')
                text = _safe_text(value)
                if text:
                    return text
        return self.default_org_name or 'Organisation'

    def _best_raw_value(self, records: list[Any], keys: list[str]) -> str:
        for record in records:
            raw = getattr(record, 'raw', {}) or {}
            for key in keys:
                text = _safe_text(raw.get(key))
                if text:
                    return text
        return ''

    def _canonical_country(self, country_raw: str) -> str:
        text = _safe_text(country_raw)
        return COUNTRY_ALIASES.get(text.lower(), text or 'Niger')

    def _display_country(self, canonical: str, raw: str, lang: str) -> str:
        if lang == 'fr':
            return COUNTRY_DISPLAY_FR.get(canonical, _safe_text(raw, canonical))
        return canonical or raw

    def _select_primary_record(self, records: list[Any]) -> Any:
        def score(record: Any) -> tuple[int, int, int, str]:
            raw = getattr(record, 'raw', {}) or {}
            media_count = len([p for p in getattr(record, 'media_files', []) if Path(p).suffix.lower() in IMAGE_EXTENSIONS])
            gps = 1 if raw.get('latitude') is not None and raw.get('longitude') is not None else 0
            desc_len = len(_safe_text(getattr(record, 'description', '') or raw.get('description') or raw.get('summary')))
            date = _safe_text(raw.get('createdAt') or getattr(record, 'created_at', ''))
            return (media_count, gps, desc_len, date)
        return sorted(records, key=score, reverse=True)[0]

    def _extract_location(self, primary_record: Any, records: list[Any], country_display: str) -> str:
        raw = getattr(primary_record, 'raw', {}) or {}
        primary_values = [raw.get('locationLabel'), getattr(primary_record, 'location', '')] if getattr(primary_record, 'kind', '') == 'story' else [getattr(primary_record, 'location', ''), raw.get('locationLabel')]
        for value in primary_values:
            text = _safe_text(value)
            if text and text != '-':
                if country_display and country_display.lower() not in text.lower():
                    return f'{text}, {country_display}'
                return text
        for record in records:
            raw = getattr(record, 'raw', {}) or {}
            text = _safe_text(raw.get('locationLabel') if getattr(record, 'kind', '') == 'story' else (getattr(record, 'location', '') or raw.get('locationLabel')))
            if text and text != '-':
                return f'{text}, {country_display}' if country_display and country_display.lower() not in text.lower() else text
        return country_display

    def _extract_activity_output(self, primary_record: Any, records: list[Any], lang: str) -> str:
        keys = [
            'output', 'outputName', 'output_name', 'outputTitle', 'output_title', 'outputLabel', 'output_label',
            'result', 'resultName', 'result_name', 'resultStatement', 'result_statement',
            'logicalFrameworkOutput', 'logframeOutput', 'projectOutput', 'project_output',
        ]
        sources: list[dict[str, Any]] = []
        for record in [primary_record, *records]:
            raw = getattr(record, 'raw', {}) or {}
            if isinstance(raw, dict):
                sources.append(raw)
                project = raw.get('project')
                if isinstance(project, dict):
                    sources.append(project)
        for source in sources:
            for key in keys:
                value = source.get(key)
                if isinstance(value, dict):
                    value = value.get('name') or value.get('title') or value.get('label') or value.get('description')
                text = self._polish_source_description(_safe_text(value))
                if text and text != '-':
                    return text
        return ''

    def _extract_activity(self, primary_record: Any, records: list[Any], lang: str) -> str:
        for value in [getattr(primary_record, 'raw', {}).get('activity'), getattr(primary_record, 'title', '')]:
            text = _clean_spaces(value)
            if text and text != '-':
                return text[:1].upper() + text[1:]
        counts = Counter(_safe_text(getattr(record, 'raw', {}).get('activity')) for record in records if _safe_text(getattr(record, 'raw', {}).get('activity')))
        if counts:
            return counts.most_common(1)[0][0]
        return 'Activité terrain' if lang == 'fr' else 'Field activity'

    def _polish_activity(self, activity: str, lang: str) -> str:
        text = _clean_spaces(activity)
        if lang != 'fr':
            return text[:1].upper() + text[1:] if text else 'Field activity'
        replacements = [
            (r'\bkits agricole\b', 'kits agricoles'),
            (r'\bkit agricole\b', 'kit agricole'),
            (r'\bgroupement agricole\b', 'groupement agricole'),
            (r'\bgroupements agricole\b', 'groupements agricoles'),
            (r'\bbénéficiaire\b', 'bénéficiaire'),
        ]
        for pattern, repl in replacements:
            text = re.sub(pattern, repl, text, flags=re.IGNORECASE)
        return text[:1].upper() + text[1:] if text else 'Activité terrain'

    def _polish_source_description(self, description: str) -> str:
        text = _clean_spaces(description)
        if not text:
            return ''
        text = re.sub(r'(\d)([A-Za-zÀ-ÿ])', r' ', text)
        text = re.sub(r'([A-Za-zÀ-ÿ])(\d)', r' ', text)
        text = re.sub(r'a\s+(\d)', r'à ', text, flags=re.IGNORECASE)
        fixes = [
            (r'kits agricole', 'kits agricoles'),
            (r'kit agricole', 'kit agricole'),
            (r'groupement et formations', 'groupements et formations'),
            (r'groupement agricole', 'groupement agricole'),
            (r'groupements agricole', 'groupements agricoles'),
            (r'agriculteur au sein', 'agriculteurs au sein'),
            (r'bénéficiaire issus', 'bénéficiaires issus'),
            (r'beneficiaire', 'bénéficiaire'),
            (r'beneficiaires', 'bénéficiaires'),
        ]
        for pattern, repl in fixes:
            text = re.sub(pattern, repl, text, flags=re.IGNORECASE)
        return text[:1].upper() + text[1:]

    def _report_title(self, activity: str, location: str, country: str, lang: str) -> str:
        short_location = location
        if ',' in short_location:
            parts = [part.strip() for part in short_location.split(',') if part.strip()]
            short_location = ', '.join(parts[:2])
        if lang == 'fr':
            return f"Rapport d’activité : {activity} à {short_location}"
        return f'Activity report: {activity} in {short_location}'

    def _detect_sectors(self, records: list[Any], activity: str, project_name: str, lang: str) -> list[dict[str, str]]:
        corpus = ' '.join([
            activity,
            project_name,
            *[_safe_text(getattr(record, 'title', '')) for record in records],
            *[_safe_text(getattr(record, 'description', '')) for record in records],
            *[_safe_text((getattr(record, 'raw', {}) or {}).get('activity')) for record in records],
            *[_safe_text((getattr(record, 'raw', {}) or {}).get('output')) for record in records],
        ]).lower()
        scored: list[tuple[int, int, dict[str, str]]] = []
        for order, rule in enumerate(SECTOR_RULES):
            score = sum(1 for keyword in rule['keywords'] if keyword.lower() in corpus)
            if score:
                scored.append((score, -order, {'key': rule['key'], 'label': rule[lang], 'icon': rule['icon']}))
        if not scored:
            return [{'key': 'coordination', 'label': 'Coordination', 'icon': 'coordination'}]
        scored.sort(key=lambda item: (-item[0], -item[1]))
        result: list[dict[str, str]] = []
        seen: set[str] = set()
        for _, __, sector in scored:
            if sector['key'] in seen:
                continue
            result.append(sector)
            seen.add(sector['key'])
            if len(result) == 2:
                break
        return result

    def _extract_indicators(self, records: list[Any], lang: str) -> list[dict[str, str]]:
        corpus = ' '.join([
            _safe_text(getattr(record, 'description', '')) + ' ' +
            _safe_text(getattr(record, 'title', '')) + ' ' +
            _safe_text((getattr(record, 'raw', {}) or {}).get('activity'))
            for record in records
        ])
        indicators: list[dict[str, str]] = []
        seen: set[tuple[str, str]] = set()
        for key, pattern, fr_label, en_label, icon, priority in INDICATOR_PATTERNS:
            for match in re.finditer(pattern, corpus, flags=re.IGNORECASE):
                number = self._normalize_number(match.group(1))
                if not number:
                    continue
                sig = (key, number)
                if sig in seen:
                    continue
                indicators.append({
                    'kind': key,
                    'value': number,
                    'label': fr_label if lang == 'fr' else en_label,
                    'icon': icon,
                    'priority': str(priority),
                })
                seen.add(sig)
        indicators.sort(key=lambda item: (-int(item.get('priority', '0')), item['label']))
        return indicators

    def _extract_primary_quantity(self, records: list[Any], sectors: list[dict[str, str]], lang: str) -> dict[str, str]:
        indicators = self._extract_indicators(records, lang)
        return indicators[0] if indicators else {'value': str(len([r for r in records if getattr(r, 'kind', '') == 'evidence']) or len(records)), 'label': 'Preuves' if lang == 'fr' else 'Evidence items', 'icon': 'check', 'kind': 'evidence'}

    def _normalize_number(self, value: str) -> str:
        cleaned = re.sub(r'[^0-9]', '', value)
        if not cleaned:
            return value.strip()
        try:
            return f'{int(cleaned):,}'.replace(',', ' ')
        except Exception:
            return cleaned

    def _all_media(self, records: list[Any]) -> list[Path]:
        media: list[Path] = []
        for record in records:
            media.extend([Path(p) for p in getattr(record, 'media_files', [])])
        return media

    def _select_hero_image(self, image_files: list[Path], primary_record: Any, records: list[Any]) -> Path | None:
        if not image_files:
            return None
        primary_images = [Path(p) for p in getattr(primary_record, 'media_files', []) if Path(p).suffix.lower() in IMAGE_EXTENSIONS]
        candidates = primary_images or image_files
        best = None
        best_score = -1.0
        for path in candidates:
            try:
                im = Image.open(path).convert('RGB')
                w, h = im.size
                small = im.resize((96, 96), Image.Resampling.LANCZOS).convert('L')
                stat = ImageStat.Stat(small)
                contrast = stat.stddev[0]
                mean = stat.mean[0]
                white_penalty = 200 if mean > 235 and contrast < 35 else 0
                aspect_bonus = 70 if w >= h else 0
                score = min(w * h / 50000, 120) + contrast + aspect_bonus - white_penalty
                if score > best_score:
                    best = path
                    best_score = score
            except Exception:
                continue
        return best or candidates[0]

    def _extract_gps(self, primary_record: Any, records: list[Any]) -> tuple[float, float, str] | None:
        candidates = [primary_record, *records]
        for record in candidates:
            raw = getattr(record, 'raw', {}) or {}
            lat = _parse_float(raw.get('latitude') or raw.get('lat'))
            lon = _parse_float(raw.get('longitude') or raw.get('lng') or raw.get('lon'))
            if lat is not None and lon is not None:
                label = _safe_text(raw.get('locationLabel') or getattr(record, 'location', ''), 'Localisation')
                return (lat, lon, label)
        return None

    def _target_group_label(self, quantity: dict[str, str], sectors: list[dict[str, str]], description: str, lang: str) -> str:
        label = quantity['label'].lower()
        corpus = f"{label} {description}".lower()
        if lang == 'fr':
            if 'agriculteur' in corpus or any(s['key'] == 'agriculture' for s in sectors):
                return 'Agriculteurs et groupements ciblés'
            if 'abri' in corpus or any(s['key'] == 'shelter' for s in sectors):
                return 'Ménages et communautés ciblées'
            if 'école' in corpus or any(s['key'] == 'education' for s in sectors):
                return 'Élèves, enseignants et communautés scolaires'
            if 'bénéficiaire' in corpus or 'participant' in corpus:
                return 'Bénéficiaires de l’activité'
            return 'Communautés ciblées'
        if 'farmer' in corpus or any(s['key'] == 'agriculture' for s in sectors):
            return 'Farmers and targeted groups'
        if 'shelter' in corpus or any(s['key'] == 'shelter' for s in sectors):
            return 'Beneficiary households and communities'
        return 'Targeted communities'

    def _evidence_type_label(self, primary_record: Any, lang: str) -> str:
        subtype = _safe_text(getattr(primary_record, 'subtype', '') or getattr(primary_record, 'raw', {}).get('type')).lower()
        if 'photo' in subtype or 'image' in subtype:
            return 'Preuve photo' if lang == 'fr' else 'Photo evidence'
        if 'video' in subtype or 'vidéo' in subtype:
            return 'Preuve vidéo' if lang == 'fr' else 'Video evidence'
        if 'document' in subtype or 'doc' in subtype:
            return 'Document source' if lang == 'fr' else 'Source document'
        return 'Preuve terrain' if lang == 'fr' else 'Field evidence'

    def _story_subject(self, primary_record: Any, lang: str) -> str:
        raw = getattr(primary_record, 'raw', {}) or {}
        for value in [raw.get('beneficiaryAlias'), getattr(primary_record, 'title', ''), raw.get('title'), raw.get('activity')]:
            text = _safe_text(value)
            if text and text != '-':
                return text
        return 'Bénéficiaire' if lang == 'fr' else 'Beneficiary'

    def _story_report_title(self, primary_record: Any, activity: str, location: str, country: str, lang: str) -> str:
        subject = self._story_subject(primary_record, lang)
        short_location = location
        if ',' in short_location:
            parts = [part.strip() for part in short_location.split(',') if part.strip()]
            short_location = ', '.join(parts[:2])
        if lang == 'fr':
            return f"Success story : {subject} à {short_location}"
        return f"Success story: {subject} in {short_location}"

    def _story_indicators(self, primary_record: Any, records: list[Any], indicators: list[dict[str, str]], lang: str) -> list[dict[str, str]]:
        result: list[dict[str, str]] = []
        seen: set[tuple[str, str]] = set()
        for item in indicators:
            sig = (_safe_text(item.get('kind')), _safe_text(item.get('value')))
            if sig in seen:
                continue
            result.append(item)
            seen.add(sig)
            if len(result) == 1:
                break
        media_count = sum(len(getattr(record, 'media_files', []) or []) for record in records)
        result.append({'kind': 'media', 'value': str(media_count), 'label': 'Médias liés' if lang == 'fr' else 'Linked media', 'icon': 'media'})
        consent = any(bool((getattr(record, 'raw', {}) or {}).get('consentGiven')) for record in records)
        result.append({'kind': 'story', 'value': str(len(records) or 1), 'label': 'Success story' if lang == 'fr' else 'Success story', 'icon': 'check'})
        if len(result) < 3:
            consent_value = ('Oui' if consent else 'Non') if lang == 'fr' else ('Yes' if consent else 'No')
            result.append({'kind': 'consent', 'value': consent_value, 'label': 'Consentement' if lang == 'fr' else 'Consent', 'icon': 'check'})
        return result[:3]

    def _story_target_group(self, primary_record: Any, description: str, lang: str) -> str:
        raw = getattr(primary_record, 'raw', {}) or {}
        beneficiary = _safe_text(raw.get('beneficiaryAlias') or getattr(primary_record, 'title', ''))
        if beneficiary:
            return beneficiary
        return 'Bénéficiaire mis en lumière' if lang == 'fr' else 'Featured beneficiary'

    def _story_narrative(self, primary_record: Any, project_name: str, activity: str, description: str, location: str, lang: str) -> str:
        raw = getattr(primary_record, 'raw', {}) or {}
        beneficiary = self._story_subject(primary_record, lang)
        summary = self._polish_source_description(_safe_text(raw.get('summary') or description))
        quote = self._polish_source_description(_safe_text(raw.get('quote')))
        if lang == 'fr':
            parts = [f"Cette success story met en lumière {beneficiary} dans le cadre du projet {project_name} à {location}."]
            if summary:
                parts.append(summary)
            elif activity:
                parts.append(f"Le récit est lié à l’activité {activity[:1].lower() + activity[1:]}.")
            if quote:
                parts.append(f"Citation clé : « {quote} »")
            parts.append("Elle apporte un éclairage qualitatif sur les changements observés, les effets perçus par le bénéficiaire et la valeur humaine de l’intervention.")
            return ' '.join(_sentence(p) for p in parts)
        parts = [f"This success story highlights {beneficiary} under the project {project_name} in {location}."]
        if summary:
            parts.append(summary)
        if quote:
            parts.append(f"Key quote: \"{quote}\"")
        parts.append("It provides qualitative insight into perceived change, beneficiary experience, and the human value of the intervention.")
        return ' '.join(_sentence(p) for p in parts)

    def _story_highlights(self, primary_record: Any, project_code: str, activity: str, location: str, media_count: int, lang: str) -> list[str]:
        raw = getattr(primary_record, 'raw', {}) or {}
        beneficiary = self._story_subject(primary_record, lang)
        consent = bool(raw.get('consentGiven'))
        if lang == 'fr':
            highlights = [f"Une success story centrée sur {beneficiary} a été documentée à {location} dans le cadre du projet {project_code}."]
            if activity:
                highlights.append(f"Le récit est rattaché à l’activité {activity[:1].lower() + activity[1:]}." )
            if media_count > 0:
                highlights.append(f"{media_count} média(s) lié(s) renforcent l’illustration visuelle de cette success story.")
            highlights.append('Le consentement à la réutilisation est confirmé.' if consent else 'Le consentement à la réutilisation doit être vérifié avant diffusion externe.')
            return [_sentence(item) for item in highlights]
        highlights = [f"A success story focusing on {beneficiary} was documented in {location} under project {project_code}."]
        if activity:
            highlights.append(f"The story is linked to the activity {activity}.")
        if media_count > 0:
            highlights.append(f"{media_count} linked media asset(s) strengthen the visual documentation of this success story.")
        highlights.append('Consent for reuse is confirmed.' if consent else 'Consent for reuse should be checked before external dissemination.')
        return [_sentence(item) for item in highlights]

    def _indicator_value(self, indicators: list[dict[str, str]], kind: str) -> str:
        for item in indicators:
            if item.get('kind') == kind:
                return item.get('value', '')
        return ''

    def _indicator_phrase(self, indicators: list[dict[str, str]], lang: str) -> str:
        focus = [item for item in indicators if item.get('kind') not in {'media', 'evidence'}][:3]
        if not focus:
            return ''
        return ', '.join([f"{item['value']} {item['label'].lower()}" for item in focus[:-1]]) + ((' et ' if lang == 'fr' else ' and ') + f"{focus[-1]['value']} {focus[-1]['label'].lower()}" if len(focus) > 1 else f"{focus[0]['value']} {focus[0]['label'].lower()}")

    def _narrative(self, activity: str, project_name: str, description: str, indicators: list[dict[str, str]], location: str, sectors: list[dict[str, str]], lang: str) -> str:
        act = activity[:1].lower() + activity[1:] if activity else ('activité terrain' if lang == 'fr' else 'field activity')
        sector_keys = {s.get('key') for s in sectors}
        beneficiaries = self._indicator_value(indicators, 'beneficiary')
        kits = self._indicator_value(indicators, 'kit')
        groups = self._indicator_value(indicators, 'group')
        shelters = self._indicator_value(indicators, 'shelter')
        phrase = self._indicator_phrase(indicators, lang)
        if lang == 'fr':
            if {'food_security', 'agriculture'} & sector_keys and (kits or beneficiaries):
                parts = [f"Cette activité de {act} documentée à {location} soutient directement les moyens d’existence et la sécurité alimentaire des ménages ciblés."]
                if phrase:
                    parts.append(f"Les données disponibles font ressortir {phrase}, offrant une lecture claire du volume d’appui réalisé.")
                if groups:
                    parts.append(f"L’organisation en {groups} groupements renforce l’entraide, la diffusion des pratiques et la pérennité de l’investissement.")
                parts.append("Le suivi de l’utilisation des kits et de la production permettra d’éclairer les décisions de réinvestissement et de mise à l’échelle.")
                return ' '.join(_sentence(p) for p in parts)
            if 'shelter' in sector_keys or shelters:
                parts = [f"Cette activité de {act} documentée à {location} contribue à améliorer les conditions de vie et de protection des ménages ciblés."]
                if phrase:
                    parts.append(f"Les données disponibles mettent en avant {phrase}, donnant une indication concrète du volume d’appui fourni.")
                parts.append("Un suivi de l’occupation, de la qualité des abris et de la satisfaction des ménages permettra de confirmer la contribution de l’intervention à la sécurité et à la dignité des bénéficiaires.")
                return ' '.join(_sentence(p) for p in parts)
            parts = [f"Cette activité de {act} a été documentée à {location} dans le cadre du projet {project_name}."]
            if phrase:
                parts.append(f"Les informations disponibles mettent en avant {phrase} comme résultats immédiatement visibles de l’activité.")
            if description:
                parts.append(f"La description source précise : {description}.")
            parts.append("Le suivi des effets observés permettra de relier plus finement les réalisations documentées aux changements attendus pour les communautés ciblées.")
            return ' '.join(_sentence(p) for p in parts)
        # English fallback
        parts = [f"This {act} activity was documented in {location} under the project {project_name}."]
        if phrase:
            parts.append(f"Available information highlights {phrase} as the most visible immediate results of the activity.")
        parts.append('Follow-up on use and observed effects will help connect documented outputs with expected changes for targeted communities.')
        return ' '.join(_sentence(p) for p in parts)

    def _highlights(self, activity: str, project_code: str, description: str, indicators: list[dict[str, str]], media_count: int, location: str, sectors: list[dict[str, str]], lang: str) -> list[str]:
        act = activity[:1].lower() + activity[1:] if activity else ('activité terrain' if lang == 'fr' else 'field activity')
        phrase = self._indicator_phrase(indicators, lang)
        sector_text = ', '.join([s['label'] for s in sectors[:2]])
        if lang == 'fr':
            highlights = [f"Une activité de {act} a été documentée à {location} dans le cadre du projet {project_code}."]
            if phrase:
                highlights.append(f"Les données source font ressortir {phrase} comme principaux chiffres opérationnels.")
            if media_count > 0:
                highlights.append(f"{media_count} média(s) lié(s) apportent un appui visuel à la documentation de l’activité.")
            if sector_text:
                highlights.append(f"L’activité est alignée principalement avec les secteurs {sector_text.lower()}.")
            return [_sentence(item) for item in highlights]
        highlights = [f"A {act} activity was documented in {location} under project {project_code}."]
        if phrase:
            highlights.append(f"Source data highlights {phrase} as the main operational figures.")
        if media_count > 0:
            highlights.append(f"{media_count} linked media asset(s) provide visual support to document the activity.")
        if sector_text:
            highlights.append(f"The activity is primarily aligned with the {sector_text} sectors.")
        return [_sentence(item) for item in highlights]

    # ------------------------------------------------------------------ map helpers
    # ------------------------------------------------------------------ map helpers

    def _shape_file(self) -> Path:
        return resource_path('desktop_sync_app/assets/maps/naturalearth_lowres/naturalearth_lowres.shp')

    def _iter_shapes(self):
        if self._shape_cache is not None:
            return self._shape_cache
        if not self._shape_file().exists():
            self._shape_cache = []
            return self._shape_cache
        reader = shapefile.Reader(str(self._shape_file()), encoding='latin1')
        self._shape_cache = list(reader.iterShapeRecords())
        return self._shape_cache

    def _country_shape(self, country: str):
        target = self._shape_country_name(country)
        for sr in self._iter_shapes():
            if self._record_name(sr.record) == target:
                return sr.shape
        return None

    def _record_name(self, record) -> str:
        try:
            return _safe_text(record['name'])
        except Exception:
            try:
                return _safe_text(record[2])
            except Exception:
                return ''

    def _shape_country_name(self, country: str) -> str:
        return COUNTRY_ALIASES.get(_safe_text(country).lower(), _safe_text(country))

    def _shape_polygons(self, shape, country: str) -> tuple[list[list[tuple[float, float]]], tuple[float, float, float, float]]:
        points = shape.points
        parts = list(shape.parts) + [len(points)]
        polygons: list[list[tuple[float, float]]] = []
        for index in range(len(parts) - 1):
            pts = points[parts[index]:parts[index + 1]]
            if len(pts) >= 3:
                poly = [(float(x), float(y)) for x, y in pts]
                if self._shape_country_name(country) == 'France':
                    xs = [p[0] for p in poly]
                    ys = [p[1] for p in poly]
                    cx, cy = sum(xs) / len(xs), sum(ys) / len(ys)
                    if not (-6.5 <= cx <= 10.5 and 41.0 <= cy <= 52.5):
                        continue
                polygons.append(poly)
        if not polygons:
            return [], tuple(float(v) for v in shape.bbox)  # type: ignore[return-value]
        xs = [x for poly in polygons for x, _ in poly]
        ys = [y for poly in polygons for _, y in poly]
        return polygons, (min(xs), min(ys), max(xs), max(ys))

    def _map_transform(self, bbox: tuple[float, float, float, float], width: int, height: int, pad: int):
        minx, miny, maxx, maxy = bbox
        scale = min((width - pad * 2) / max(maxx - minx, 0.001), (height - pad * 2) / max(maxy - miny, 0.001))
        offset_x = (width - (maxx - minx) * scale) / 2
        offset_y = (height - (maxy - miny) * scale) / 2

        def mapper(lon: float, lat: float) -> tuple[float, float]:
            x = offset_x + (lon - minx) * scale
            y = offset_y + (maxy - lat) * scale
            return x, y
        return mapper

    def _point_in_bbox(self, lon: float, lat: float, bbox: tuple[float, float, float, float]) -> bool:
        return bbox[0] <= lon <= bbox[2] and bbox[1] <= lat <= bbox[3]

    # ------------------------------------------------------------------ docx formatting helpers

    def _spacer(self, doc: Document, points: float) -> None:
        if points <= 0:
            return
        p = doc.add_paragraph()
        self._pconf(p, before=0, after=0, line=0.1)
        p.add_run('\u00A0').font.size = Pt(points)

    def _spacer_cell(self, cell, points: float) -> None:
        p = cell.add_paragraph()
        self._pconf(p, before=0, after=0, line=0.1)
        p.add_run('\u00A0').font.size = Pt(points)

    def _section_title(self, cell, title: str, icon: str | None = None, line_chars: str = '──────────', underline_width: float = 2.2) -> None:
        p = cell.add_paragraph()
        self._pconf(p, before=0, after=0, line=0.98)
        if icon:
            p.add_run().add_picture(str(self._plain_icon(icon)), width=Inches(0.31))
            self._run(p, '  ', size=1)
        self._run(p, title, bold=True, size=12.6, color=BLUE)
        p = cell.add_paragraph()
        self._pconf(p, before=0, after=6, line=0.56)
        self._run(p, '─' * max(11, int(underline_width * 10)), size=5.6, color=ORANGE)

    def _cell_p(self, cell):
        return cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

    def _run(self, paragraph, text: str, bold: bool = False, size: float = 8.5, color: str = DARK, italic: bool = False, font: str = 'Arial'):
        text_str = str(text)
        visible_text = text_str.strip()
        if visible_text and any(ch.isalnum() for ch in visible_text):
            size = max(10.0, float(size))
        run = paragraph.add_run(text_str)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = _rgb(color)
        run.font.name = font
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
        except Exception:
            pass
        return run

    def _pconf(self, paragraph, before: float = 0, after: float = 0, line: float = 1.0, align=None) -> None:
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)
        paragraph.paragraph_format.line_spacing = line
        if align is not None:
            paragraph.alignment = align

    def _shade_cell(self, cell, fill: str) -> None:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.find(qn('w:shd'))
        if shd is None:
            shd = OxmlElement('w:shd')
            tcPr.append(shd)
        shd.set(qn('w:fill'), fill)

    def _border_cell(self, cell, color: str = BORDER, size: int = 4) -> None:
        tcPr = cell._tc.get_or_add_tcPr()
        borders = tcPr.first_child_found_in('w:tcBorders')
        if borders is None:
            borders = OxmlElement('w:tcBorders')
            tcPr.append(borders)
        for edge in ('top', 'left', 'bottom', 'right'):
            el = borders.find(qn(f'w:{edge}'))
            if el is None:
                el = OxmlElement(f'w:{edge}')
                borders.append(el)
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), str(size))
            el.set(qn('w:color'), color)

    def _table_no_borders(self, table) -> None:
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                borders = tcPr.first_child_found_in('w:tcBorders')
                if borders is None:
                    borders = OxmlElement('w:tcBorders')
                    tcPr.append(borders)
                for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                    el = borders.find(qn(f'w:{edge}'))
                    if el is None:
                        el = OxmlElement(f'w:{edge}')
                        borders.append(el)
                    el.set(qn('w:val'), 'nil')

    def _set_table_grid(self, table, widths: list[float]) -> None:
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl.insert(0, tbl_pr)
        total_twips = int(sum(widths) * 1440)

        # LibreOffice/Word may otherwise keep an implicit table indentation even when
        # cell widths are correct. Explicit table width is critical here: the header
        # must be full-bleed, while body tables use a centered 2cm-safe width.
        tbl_w = tbl_pr.find(qn('w:tblW'))
        if tbl_w is None:
            tbl_w = OxmlElement('w:tblW')
            tbl_pr.append(tbl_w)
        tbl_w.set(qn('w:w'), str(total_twips))
        tbl_w.set(qn('w:type'), 'dxa')

        tbl_ind = tbl_pr.find(qn('w:tblInd'))
        if tbl_ind is None:
            tbl_ind = OxmlElement('w:tblInd')
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn('w:w'), '0')
        tbl_ind.set(qn('w:type'), 'dxa')

        old_grid = tbl.tblGrid
        if old_grid is not None:
            tbl.remove(old_grid)
        grid = OxmlElement('w:tblGrid')
        for width in widths:
            col = OxmlElement('w:gridCol')
            col.set(qn('w:w'), str(int(width * 1440)))
            grid.append(col)
        # Keep OOXML child order valid: tblPr must precede tblGrid.
        try:
            insert_at = list(tbl).index(tbl_pr) + 1
        except ValueError:
            insert_at = 0
        tbl.insert(insert_at, grid)
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    self._set_cell_width(row.cells[idx], Inches(width))

    def _set_table_indent(self, table, twips: int) -> None:
        tbl_pr = table._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            table._tbl.insert(0, tbl_pr)
        tbl_ind = tbl_pr.find(qn('w:tblInd'))
        if tbl_ind is None:
            tbl_ind = OxmlElement('w:tblInd')
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn('w:w'), str(int(twips)))
        tbl_ind.set(qn('w:type'), 'dxa')

    def _set_cell_width(self, cell, width) -> None:
        twips = int(width.inches * 1440) if hasattr(width, 'inches') else int(width)
        tcPr = cell._tc.get_or_add_tcPr()
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is None:
            tcW = OxmlElement('w:tcW')
            tcPr.append(tcW)
        tcW.set(qn('w:w'), str(twips))
        tcW.set(qn('w:type'), 'dxa')

    def _set_row_height(self, row, height) -> None:
        row.height = height
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    def _cell_margins(self, cell, top=70, bottom=70, start=70, end=70) -> None:
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = tcPr.first_child_found_in('w:tcMar')
        if tcMar is None:
            tcMar = OxmlElement('w:tcMar')
            tcPr.append(tcMar)
        for name, val in [('top', top), ('bottom', bottom), ('start', start), ('end', end)]:
            node = tcMar.find(qn(f'w:{name}'))
            if node is None:
                node = OxmlElement(f'w:{name}')
                tcMar.append(node)
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')

    # ------------------------------------------------------------------ PIL drawing helpers

    def _draw_line_icon_pil(self, draw: ImageDraw.ImageDraw, icon: str, center: tuple[int, int], color: tuple[int, int, int, int], scale: float = 1.0) -> None:
        x, y = center
        s = 12 * scale
        w = max(4, int(1.2 * scale))
        c = color
        if icon in {'doc', 'document'}:
            draw.rounded_rectangle((x - s*1.0, y - s*1.35, x + s*1.0, y + s*1.35), radius=int(s*0.18), outline=c, width=w)
            draw.line((x + s*0.45, y - s*1.35, x + s*1.0, y - s*0.80, x + s*1.0, y - s*1.35), fill=c, width=w)
            for yy in (-0.55, -0.15, 0.25, 0.65):
                draw.line((x - s*0.55, y + s*yy, x + s*0.45, y + s*yy), fill=c, width=w)
        elif icon in {'bar', 'barchart'}:
            draw.line((x - s*1.2, y + s*1.1, x + s*1.25, y + s*1.1), fill=c, width=w)
            draw.rounded_rectangle((x - s*0.95, y + s*0.20, x - s*0.45, y + s*1.05), radius=int(s*0.08), fill=c)
            draw.rounded_rectangle((x - s*0.20, y - s*0.35, x + s*0.30, y + s*1.05), radius=int(s*0.08), fill=c)
            draw.rounded_rectangle((x + s*0.55, y - s*0.85, x + s*1.05, y + s*1.05), radius=int(s*0.08), fill=c)
        elif icon in {'people', 'person'}:
            draw.ellipse((x - s*0.35, y - s*1.05, x + s*0.35, y - s*0.35), outline=c, width=w)
            draw.arc((x - s*0.95, y - s*0.10, x + s*0.95, y + s*1.20), 200, 340, fill=c, width=w)
            draw.ellipse((x - s*1.18, y - s*0.80, x - s*0.68, y - s*0.30), outline=c, width=max(3, w-1))
            draw.arc((x - s*1.55, y - s*0.10, x - s*0.40, y + s*0.95), 215, 340, fill=c, width=max(3, w-1))
            draw.ellipse((x + s*0.68, y - s*0.80, x + s*1.18, y - s*0.30), outline=c, width=max(3, w-1))
            draw.arc((x + s*0.40, y - s*0.10, x + s*1.55, y + s*0.95), 200, 325, fill=c, width=max(3, w-1))
        elif icon == 'org':
            draw.rounded_rectangle((x - s*1.05, y - s*1.1, x + s*1.05, y + s*1.1), radius=int(s*0.12), outline=c, width=w)
            for dx in (-0.55, 0, 0.55):
                draw.line((x + s*dx, y - s*0.95, x + s*dx, y + s*0.95), fill=c, width=w)
            draw.line((x - s*1.15, y - s*1.15, x + s*1.15, y - s*1.15), fill=c, width=w)
            draw.line((x - s*1.15, y + s*1.15, x + s*1.15, y + s*1.15), fill=c, width=w)
        elif icon in {'media', 'camera'}:
            draw.rounded_rectangle((x - s*1.20, y - s*0.95, x + s*1.20, y + s*0.95), radius=int(s*0.15), outline=c, width=w)
            draw.rectangle((x - s*0.82, y - s*0.55, x + s*0.40, y + s*0.50), outline=c, width=max(3, w-1))
            if icon == 'camera':
                draw.ellipse((x - s*0.25, y - s*0.20, x + s*0.15, y + s*0.20), outline=c, width=max(3, w-1))
                draw.rectangle((x - s*0.72, y - s*0.75, x - s*0.25, y - s*0.45), fill=c)
            else:
                draw.polygon([(x - s*0.05, y - s*0.18), (x - s*0.05, y + s*0.18), (x + s*0.35, y)], fill=c)
                draw.line((x - s*1.20, y - s*0.55, x - s*1.48, y - s*0.35, x - s*1.20, y - s*0.15), fill=c, width=w)
        elif icon in {'check'}:
            draw.rounded_rectangle((x - s*0.95, y - s*1.20, x + s*0.95, y + s*1.15), radius=int(s*0.15), outline=c, width=w)
            draw.rectangle((x - s*0.40, y - s*1.45, x + s*0.40, y - s*1.10), outline=c, width=max(3, w-1))
            draw.line((x - s*0.45, y + s*0.10, x - s*0.12, y + s*0.45, x + s*0.52, y - s*0.30), fill=c, width=w)
        elif icon in {'calendar'}:
            draw.rounded_rectangle((x - s*1.0, y - s*1.10, x + s*1.0, y + s*1.0), radius=int(s*0.15), outline=c, width=w)
            draw.line((x - s*1.0, y - s*0.45, x + s*1.0, y - s*0.45), fill=c, width=w)
            draw.line((x - s*0.45, y - s*1.35, x - s*0.45, y - s*0.75), fill=c, width=w)
            draw.line((x + s*0.45, y - s*1.35, x + s*0.45, y - s*0.75), fill=c, width=w)
            for dx in (-0.42, 0.0, 0.42):
                for dy in (0.0, 0.45):
                    draw.ellipse((x + s*dx - s*0.06, y + s*dy - s*0.06, x + s*dx + s*0.06, y + s*dy + s*0.06), fill=c)
        elif icon in {'pin', 'pin_round'}:
            draw.ellipse((x - s*0.62, y - s*1.18, x + s*0.62, y + s*0.02), outline=c, width=w)
            draw.polygon([(x, y + s*1.20), (x - s*0.42, y - s*0.05), (x + s*0.42, y - s*0.05)], outline=c, width=w)
            draw.ellipse((x - s*0.18, y - s*0.78, x + s*0.18, y - s*0.42), outline=c, width=max(3, w-1))
        elif icon == 'star':
            pts = []
            for i in range(10):
                a = -math.pi / 2 + i * math.pi / 5
                rr = s * 1.08 if i % 2 == 0 else s * 0.46
                pts.append((x + rr * math.cos(a), y + rr * math.sin(a)))
            draw.line(pts + [pts[0]], fill=c, width=w, joint='curve')
        elif icon == 'sector':
            draw.ellipse((x - s*1.18, y - s*1.18, x + s*1.18, y + s*1.18), outline=c, width=w)
            draw.ellipse((x - s*0.55, y - s*0.55, x + s*0.55, y + s*0.55), outline=c, width=w)
            draw.line((x, y - s*1.45, x, y - s*0.80), fill=c, width=w)
            draw.line((x, y + s*0.80, x, y + s*1.45), fill=c, width=w)
            draw.line((x - s*1.45, y, x - s*0.80, y), fill=c, width=w)
            draw.line((x + s*0.80, y, x + s*1.45, y), fill=c, width=w)
            draw.ellipse((x - s*0.10, y - s*0.10, x + s*0.10, y + s*0.10), fill=c)
        else:
            draw.ellipse((x - 10, y - 10, x + 10, y + 10), fill=c)

    def _draw_sector_icon_pil(self, draw: ImageDraw.ImageDraw, icon: str, center: tuple[int, int], color: tuple[int, int, int, int], scale: float = 1.0) -> None:
        x, y = center
        s = 12 * scale
        w = max(4, int(1.15 * scale))
        c = color
        if icon == 'shelter':
            draw.line((x - s*1.15, y - s*0.05, x, y - s*1.05, x + s*1.15, y - s*0.05), fill=c, width=w, joint='curve')
            draw.rounded_rectangle((x - s*0.82, y - s*0.05, x + s*0.82, y + s*1.00), radius=int(s*0.10), outline=c, width=w)
            draw.line((x - s*0.18, y + s*1.00, x - s*0.18, y + s*0.35), fill=c, width=w)
            draw.line((x + s*0.18, y + s*1.00, x + s*0.18, y + s*0.35), fill=c, width=w)
            draw.line((x - s*0.18, y + s*0.35, x + s*0.18, y + s*0.35), fill=c, width=w)
        elif icon in {'food', 'nutrition'}:
            draw.arc((x - s*1.15, y - s*0.25, x + s*1.15, y + s*1.05), 0, 180, fill=c, width=w)
            draw.line((x - s*0.98, y + s*0.40, x + s*0.98, y + s*0.40), fill=c, width=w)
            draw.line((x - s*0.35, y - s*0.80, x - s*0.35, y - s*0.10), fill=c, width=w)
            draw.line((x - s*0.35, y - s*0.80, x - s*0.72, y - s*1.08), fill=c, width=max(3, w-1))
            draw.line((x - s*0.35, y - s*0.80, x - s*0.03, y - s*1.12), fill=c, width=max(3, w-1))
            if icon == 'nutrition':
                draw.line((x + s*0.35, y - s*0.55, x + s*0.35, y + s*0.20), fill=c, width=w)
                draw.arc((x + s*0.08, y - s*0.92, x + s*0.62, y - s*0.38), 190, 20, fill=c, width=max(3, w-1))
        elif icon == 'agriculture':
            draw.line((x, y + s*1.15, x, y - s*0.90), fill=c, width=w)
            for ox, oy, flip in [(-0.12, -0.65, -1), (0.18, -0.25, 1), (-0.15, 0.15, -1), (0.20, 0.55, 1)]:
                bx = x + s*ox
                by = y + s*oy
                if flip < 0:
                    draw.arc((bx - s*0.70, by - s*0.35, bx, by + s*0.35), 210, 35, fill=c, width=w)
                else:
                    draw.arc((bx, by - s*0.35, bx + s*0.70, by + s*0.35), 145, 330, fill=c, width=w)
        elif icon == 'wash':
            draw.line((x, y - s*1.15, x - s*0.55, y - s*0.10, x, y + s*0.85, x + s*0.55, y - s*0.10, x, y - s*1.15), fill=c, width=w, joint='curve')
            draw.arc((x - s*0.40, y - s*0.10, x + s*0.40, y + s*0.65), 200, 340, fill=c, width=max(3, w-1))
        elif icon == 'health':
            draw.rounded_rectangle((x - s*0.32, y - s*1.00, x + s*0.32, y + s*1.00), radius=int(s*0.08), fill=c)
            draw.rounded_rectangle((x - s*1.00, y - s*0.32, x + s*1.00, y + s*0.32), radius=int(s*0.08), fill=c)
        elif icon == 'education':
            draw.line((x - s*1.20, y - s*0.35, x, y - s*1.00, x + s*1.20, y - s*0.35), fill=c, width=w)
            draw.line((x, y - s*1.00, x, y + s*0.90), fill=c, width=max(3, w-1))
            draw.line((x - s*0.85, y + s*0.15, x + s*0.85, y + s*0.15), fill=c, width=w)
            draw.line((x - s*0.80, y + s*0.90, x + s*0.80, y + s*0.90), fill=c, width=w)
            draw.line((x - s*0.80, y + s*0.15, x - s*0.80, y + s*0.90), fill=c, width=w)
            draw.line((x + s*0.80, y + s*0.15, x + s*0.80, y + s*0.90), fill=c, width=w)
        elif icon == 'protection':
            pts = [(x, y - s*1.20), (x + s*0.85, y - s*0.65), (x + s*0.65, y + s*0.65), (x, y + s*1.15), (x - s*0.65, y + s*0.65), (x - s*0.85, y - s*0.65), (x, y - s*1.20)]
            draw.line(pts, fill=c, width=w, joint='curve')
        elif icon == 'cash':
            draw.rounded_rectangle((x - s*1.20, y - s*0.75, x + s*1.20, y + s*0.75), radius=int(s*0.15), outline=c, width=w)
            draw.ellipse((x - s*0.38, y - s*0.38, x + s*0.38, y + s*0.38), outline=c, width=w)
            draw.line((x, y - s*0.65, x, y + s*0.65), fill=c, width=max(3, w-1))
        else:
            self._draw_line_icon_pil(draw, 'sector', center, color, scale)

    def _draw_pin_pil(self, draw: ImageDraw.ImageDraw, x: float, y: float, size: int) -> None:
        orange = _pil_rgb(ORANGE)
        draw.ellipse((x - size * 0.55, y - size * 1.05, x + size * 0.55, y + size * 0.05), fill=orange)
        draw.polygon([(x, y + size * 0.85), (x - size * 0.42, y - size * 0.05), (x + size * 0.42, y - size * 0.05)], fill=orange)
        draw.ellipse((x - size * 0.2, y - size * 0.70, x + size * 0.2, y - size * 0.30), fill=_pil_rgb(WHITE))

    def _font(self, size: int, bold: bool = False):
        candidates = [
            r'C:\Windows\Fonts\arialbd.ttf' if bold else r'C:\Windows\Fonts\arial.ttf',
            '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf' if bold else '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        ]
        for path in candidates:
            try:
                if Path(path).exists():
                    return ImageFont.truetype(path, size=size)
            except Exception:
                continue
        return ImageFont.load_default()

    def _text_w(self, draw: ImageDraw.ImageDraw, text: str, font) -> int:
        try:
            return int(draw.textbbox((0, 0), text, font=font)[2])
        except Exception:
            return int(draw.textlength(text, font=font))

    def _center_text(self, draw: ImageDraw.ImageDraw, text: str, box: tuple[int, int, int, int], font, fill: str) -> None:
        x1, y1, x2, y2 = box
        tw = self._text_w(draw, text, font)
        th = int(draw.textbbox((0, 0), text, font=font)[3])
        draw.text((x1 + (x2 - x1 - tw) / 2, y1 + (y2 - y1 - th) / 2), text, fill=fill, font=font)

    def _truncate_text(self, text: str, max_chars: int) -> str:
        text = _clean_spaces(text)
        return text if len(text) <= max_chars else text[:max_chars - 1].rstrip() + '…'
